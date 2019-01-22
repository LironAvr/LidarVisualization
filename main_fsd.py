# -*- coding: utf-8 -*-
from config import *
import StringIO
import hashlib
import os
import random
import re
import sys
from html_decode_support import stringToConvert
import uuid
import re
import urllib2
from random import randint
from urlparse import urlparse
import time
import oauth2
import urllib2
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate, InlineImage, R
from elasticsearch import Elasticsearch
from flask import *
from flask_httpauth import HTTPBasicAuth

from extensions import *
from sqlite import *
from es import *


# Hàm phụ trợ:  Kiểm tra internet
def internet_on():
    try:
        urllib2.urlopen('http://216.58.192.142', timeout=1)
        return True
    except urllib2.URLError as err:
        return False


# Hàm phụ trợ:  cộng chuỗi nâng cao
def advance_plus_string(origi, new_string, only_one=0, sperate=', '):
    ## Kiểm tra null
    if new_string == '':
        return origi
    if origi != '' and only_one == 0:
        origi += sperate + new_string
    elif origi == '' and only_one == 0:
        origi += new_string
    elif origi != '' and only_one != 0:
        pass
    elif origi == '' and only_one != 0:
        origi += new_string
    return origi


# Hàm phụ trợ:  Do một số biến đã sử dụng list thay vì dùng dict để lưu trữ dữ liệu, ta phải cần hàm này để
# tìm chính xác các trường
def find_data(q, name):
    for i in q:
        if i[2] == name:
            return i
    return ['', '', '', '']


# Hàm phụ trợ: Một số dữ liệu thuộc dạng nhúng HTML, sử dụng hàm này để biển đổi chúng thành văn bản bình thường
def output_html(q):
    q = q.replace("</br>", "\n")
    q = q.replace("< br / >", "\n")
    q = re.sub('<[^>]*>', '', q)
    return q


# Hàm phụ trợ: Chuyển tiếng việt thành tiếng việt không dấu
def no_accent_vietnamese(utf8_str):
    # Lấy của https://gist.github.com/thuandt/3421905, có sửa lại một ít
    INTAB = "ạảãàáâậầấẩẫăắằặẳẵóòọõỏôộổỗồốơờớợởỡéèẻẹẽêếềệểễúùụủũưựữửừứíìịỉĩýỳỷỵỹđẠẢÃÀÁÂẬẦẤẨẪĂẮẰẶẲẴÓÒỌÕỎÔỘỔỖỒỐƠỜỚỢỞỠÉÈẺẸẼÊẾỀỆỂỄÚÙỤỦŨƯỰỮỬỪỨÍÌỊỈĨÝỲỶỴỸĐ"
    INTAB = [ch.encode('utf8') for ch in unicode(INTAB, 'utf8')]

    OUTTAB = "a" * 17 + "o" * 17 + "e" * 11 + "u" * 11 + "i" * 5 + "y" * 5 + "d" + \
             "A" * 17 + "O" * 17 + "E" * 11 + "U" * 11 + "I" * 5 + "Y" * 5 + "D"

    r = re.compile("|".join(INTAB))
    replaces_dict = dict(zip(INTAB, OUTTAB))
    return r.sub(lambda m: replaces_dict[m.group(0)], utf8_str)


# Hàm phụ trợ: Nhận biết liên kết có thuộc tên miền cho trước không
def detect_website_link(q):
    result = ''
    urls = re.findall('http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', q)
    for i in urls:
        o = urlparse(i)
        hostname = o.hostname.lower()
        if hostname in ['www.facebook.com', 'facebook.com', 'fb.com', 'twitter.com', 'm.facebook.com']:
            return i
    return result


# Hàm phụ trợ: tách ngày tháng sinh nhật (Chưa được dùng)
def seperate_birthday(q):
    b_day = '0'
    b_month = '0'
    b_year = '0'
    q = re.sub(r'[^0-9\s]', '', q)  # remove all anphabes
    a = q.split()
    # Full
    if len(a) == 3:
        b_day = a[0]
        b_month = a[1]
        b_year = a[2]
    # day and month only
    elif len(a) == 2:
        b_day = a[0]
        b_month = a[1]
    # year only
    elif len(a) == 1:
        b_year = a[0]
    # remove the first 0 in the day and the month
    b_day = int(b_day)
    b_month = int(b_month)
    return [str(b_day), str(b_month), b_year]


# Hàm phụ trợ: Chọn định danh
def apply_cookie(ckid):
    if int(ckid) > 0:
        a = apply_cookie_modal(ckid, session['id'])
        if a == []:
            return FB_COOKIE
        return a[0][0]
    else:
        return FB_COOKIE


# Hàm phụ trợ: Lấy kết quả từ Bing từ một website xác định, có chia trang
def bing_web_related(q, site, first):
    results = {}
    gp_web_results = []
    r = requests.get('http://www.bing.com/search?q=site:' + site + ' ' + q + '&first=' + str(first))
    soup = BeautifulSoup(r.content, "lxml")
    mydivs = soup.findAll("li", {"class": "b_algo"})
    for mydiv in mydivs:
        if mydiv.a.attrs['href'].decode('utf-8') not in gp_web_results:
            gp_web_results.append(mydiv.a.attrs['href'])
    a = {'ok': 1, 'ketqua': gp_web_results}
    results.update(a)
    return results


# Hàm phụ trợ: Kiểm tra xem có tài khoản twitter với email cho trước không
def twitter_available(email):
    results = {}
    auth = []
    a = {'ok': 1}
    results.update(a)
    r = requests.get('https://twitter.com/users/email_available?email=' + email)
    if r.content.find('Email has already been taken') != -1:
        auth.append('twitter')
    a = {'auth': auth}
    results.update(a)
    return results


# Hàm phụ trợ: bypass facebook
def CrawlFbBypass(url, key, cookie):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    data = {'_orig_post_vars': '', '_wap_notice_shown': '', 'fb_dtsg': key}
    r = requests.post(url, headers=headers, data=data)
    return r.content


# Hàm phụ trợ: Tìm group facebook (Đầu vào: uid, đầu ra:GroupName,GroupLink)
def GetFbGroup(uid, cookie):
    stop = 0
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    url = 'https://mbasic.facebook.com/search/' + uid + '/groups/intersect'
    results = []
    while stop == 0:
        r = requests.get(url, headers=headers)
        content = r.content
        # Kiểm tra chặn
        if content.find('Cảnh báo: Vui lòng chậm lại</h3>') != -1:
            # Tìm key
            start1 = content.find('fb_dtsg')
            start2 = content.find('value="', start1) + 7
            end = content.find('"', start2)
            key = content[start2:end]
            print 'bi chan roi'
            # content = CrawlFbBypass(url, key, cookie)
        if content.find('Xin lỗi, chúng tôi không thể tìm được bất kỳ kết quả nào cho bạn.</div>') != -1:
            stop = 1
        soup = BeautifulSoup(content, "lxml")
        if len(soup(id="BrowseResultsContainer")) > 0:
            a = str(soup(id="BrowseResultsContainer")[0])
            a = a.split('role="presentation">')
            for i in a:
                # Đóng tag
                i = i + ">"
                soup1 = BeautifulSoup(i, "lxml")
                # Tên group
                NameGroup = str(soup1.div)
                NameGroup = re.sub('<[^>]*>', '', NameGroup)
                if NameGroup != '':
                    results.append({'GroupName': NameGroup, 'GroupLink': soup1.a['href'].split('?__xts__')[0]})
        # Sang trang mới
        stop = 1
        if content.find('"><span>Xem thêm kết quả</span>') != -1:
            end = content.find('"><span>Xem thêm kết quả</span>')
            start = content.rfind('href="', 0, end) + len('href="')
            new_page_link = content[start:end]
            url = new_page_link
        else:
            stop = 1
    return results


# Hàm phụ trợ: Lấy tất cả comment post public
def GetCommentPost(uid, page_max, cookie, GetComment=0):
    page = 0
    stop = 0
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    url = 'https://mbasic.facebook.com/search/' + uid + '/stories-commented'
    results = []
    while page < page_max and stop == 0:
        page += 1
        r = requests.get(url, headers=headers)
        content = r.content
        # Kiểm tra chặn
        if content.find('Cảnh báo: Vui lòng chậm lại</h3>') != -1:
            # Tìm key
            start1 = content.find('fb_dtsg')
            start2 = content.find('value="', start1) + 7
            end = content.find('"', start2)
            key = content[start2:end]
            content = CrawlFbBypass(url, key, cookie)
        if content.find('Xin lỗi, chúng tôi không thể tìm được bất kỳ kết quả nào cho bạn.</div>') != -1:
            stop = 1
        soup = BeautifulSoup(content, "lxml")
        if len(soup(id="BrowseResultsContainer")) > 0:
            a = str(soup(id="BrowseResultsContainer")[0])
            a = a.split('role="presentation">')
            for i in a[-1:]:
                result = {}
                # print re.sub('<[^>]*>', '', i)
                # Lấy người đăng bài
                start = i.find("<h3 class=")
                end = i.find("</h3>")
                author = i[start:end]
                result['name'] = re.sub('<[^>]*>', '', author)
                # Lấy nội dung
                start = i.find("""data-ft='{"tn":"*s"}' style="">""")
                end = i.find("""data-ft='{"tn":"H"}'>""", start)
                content = "<" + i[start:end] + ">"
                result['value'] = re.sub('<[^>]*>', '', content)
                # Lấy thời gian đăng bài
                start = i.find("<abbr>")
                end = i.find("</abbr>", start)
                time = i[start:end]
                result['time'] = re.sub('<[^>]*>', '', time)
                # Lấy comment
                if i.find("bình luận") and GetComment == 1:
                    start = i.rfind('href="', 0, i.find("bình luận")) + 6
                    end = i.find('">', start)
                    _LinkComment = i[start:end]
                    _LinkComment = _LinkComment.replace("&amp;", "&")
                    r2 = requests.get("https://mbasic.facebook.com" + _LinkComment, headers=headers)
                    b = r2.content.split('<h3><a class=')
                    if len(b) > 0:
                        result_comment = {}
                        for j in b[1:]:
                            # Lấy tên người comment
                            end = j.find('</h3>')
                            NameComment = "<" + j[0:end]
                            result_comment['name'] = re.sub('<[^>]*>', '', NameComment)
                            # Lấy nội dung comment
                            start = j.find('<div class=')
                            end = j.find('</div>')
                            ContentComment = j[start:end]
                            result_comment['value'] = re.sub('<[^>]*>', '', ContentComment)
                            result['comments'] = result_comment

                results.append(result)
        # Sang trang mới
        if content.find('"><span>Xem thêm kết quả</span>') != -1:
            end = content.find('"><span>Xem thêm kết quả</span>')
            start = content.rfind('href="', 0, end) + len('href="')
            new_page_link = content[start:end]
            url = new_page_link
        else:
            stop = 1
    return results


# Hàm phụ trợ: Haveibeenpwned
def haveibeenpwned(q):
    results = []
    url = "https://haveibeenpwned.com/api/v2/breachedaccount/" + q
    r = requests.get(url)
    if r.status_code != 200:
        return []
    a = json.loads(r.content)
    for i in a:
        results.append({'title': i['Title'], 'domain': i['Domain']})
    return results


# Hàm phụ trợ: Lấy các địa điểm đã đến thăm
def GetPlaceVisited(uid, page_max, cookie):
    page = 0
    stop = 0
    results = []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    url = 'https://mbasic.facebook.com/search/' + uid + '/places-visited/'
    while page < page_max and stop == 0:
        page += 1
        result = {}
        r = requests.get(url, headers=headers)
        content = r.content
        # Kiểm tra chặn
        if content.find('Cảnh báo: Vui lòng chậm lại</h3>') != -1:
            # Tìm key
            start1 = content.find('fb_dtsg')
            start2 = content.find('value="', start1) + 7
            end = content.find('"', start2)
            key = content[start2:end]
            content = CrawlFbBypass(url, key, cookie)
        if content.find('Xin lỗi, chúng tôi không thể tìm được bất kỳ kết quả nào cho bạn.</div>') != -1:
            stop = 1
        soup = BeautifulSoup(content, "lxml")
        if len(soup(id="BrowseResultsContainer")) > 0:
            a = str(soup(id="BrowseResultsContainer")[0])
            a = a.split('role="presentation">')
            for i in a[1:]:
                result = {}
                # Lấy tên
                start = i.find('<div class=')
                end = i.find('</div>', start)
                name = i[start:end]
                name = re.sub('<[^>]*>', '', name)
                result['name'] = name
                # Lấy type
                start = i.find('<div class=', end)
                end = i.find('</div>', start)
                LocationType = i[start:end]
                LocationType = re.sub('<[^>]*>', '', LocationType)
                result['LocationType'] = LocationType
                results.append(result)
        # Sang trang mới
        if content.find('"><span>Xem thêm kết quả</span>') != -1:
            end = content.find('"><span>Xem thêm kết quả</span>')
            start = content.rfind('href="', 0, end) + len('href="')
            new_page_link = content[start:end]
            url = new_page_link
        else:
            stop = 1
    return results


# Hàm phụ trợ: Lấy tất cả status
def GetFbStatus(username, cookie, page_max, GetComment=0, SearchMode=0, NextURL=''):
    stop = 0
    page = 0
    resultss = {}
    results = []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    if username.find('profile.php?id=') != -1:
        url = 'https://mbasic.facebook.com/' + username + '&v=timeline'
    else:
        url = 'https://mbasic.facebook.com/' + username + '?v=timeline'
    # Nếu trong chế độ tìm kiếm thì URL khởi đầu sẽ là NextURL
    if SearchMode == 1 and NextURL != '':
        url = NextURL
    while page < page_max and stop == 0:
        page += 1
        r = requests.get(url, headers=headers)
        soup = BeautifulSoup(r.content, "lxml")
        ## Bs nó không cho tách luôn nên làm lằng ngoằng như dưới
        new_content = ""
        ## Lay noi dung trong timelinebody
        for el in soup.find_all('div', attrs={'id': 'timelineBody'}):
            new_content += str(el)
        if new_content == '':
            new_content = r.content
        a = new_content.split('>Báo cáo</a></div></div>')  # Dùng báo cáo làm phân tách
        ## Bỏ dòng cuối thừa for in -1
        for i in a[:-1]:
            result = {}
            # Lấy người đăng bài
            start = i.find("<h3 class=")
            end = i.find("</h3>")
            author = i[start:end]
            result['name'] = re.sub('<[^>]*>', '', author)
            # Lấy nội dung người đăng bài
            start = end
            end = i.find('<abbr>', start)
            content = i[start:end]
            result['value'] = re.sub('<[^>]*>', '', content.decode('utf-8', 'replace'))
            # Lấy time
            start = end
            end = i.find('</abbr>', start)
            time = i[start:end]
            result['time'] = re.sub('<[^>]*>', '', time)
            # Lấy comment
            if i.find("bình luận") and GetComment == 1:
                start = i.rfind('href="', 0, i.find("bình luận")) + 6
                end = i.find('">', start)
                _LinkComment = i[start:end]
                _LinkComment = _LinkComment.replace("&amp;", "&")
                r2 = requests.get("https://mbasic.facebook.com" + _LinkComment, headers=headers)
                b = r2.content.split('<h3><a class=')
                if len(b) > 0:
                    comment_results = []
                    for j in b[1:]:
                        comment_result = {}
                        # Lấy tên người comment
                        end = j.find('</h3>')
                        NameComment = "<" + j[0:end]
                        comment_result['name'] = re.sub('<[^>]*>', '', NameComment)
                        # Lấy nội dung comment
                        start = j.find('<div class=')
                        end = j.find('</div>')
                        ContentComment = j[start:end]
                        comment_result['value'] = re.sub('<[^>]*>', '', ContentComment.decode('utf-8', 'replace'))
                        comment_results.append(comment_result)
                    result['comments'] = comment_results
            results.append(result)
        # Sang trang mới
        if r.content.find('<span>Xem tin khác</span>') != -1 or r.content.find('Hiển thị thêm</a>') != -1:
            end = r.content.find('<span>Xem tin khác</span>') - 2
            if end == -3:
                end = r.content.find('Hiển thị thêm</a>') - 2
            start = r.content.rfind('href="', 0, end) + len('href="')
            new_page_link = r.content[start:end]
            url = 'https://mbasic.facebook.com/' + new_page_link
            url = url.replace('&amp;', '&')
        else:
            stop = 1
    resultss['NextURL'] = url
    resultss['results'] = results
    return resultss


# Hàm phụ trợ: Lấy các post hoặc comment có chứa nội dung cho trước
def SearchFbStatus(results, SearchKey, ValueName, CommentName):
    new_results = []
    if SearchKey == '*':
        return results
    # Tìm trong post
    for i in results:
        a = i[ValueName].lower()
        if a.find(SearchKey.lower()) != -1:
            new_results.append(i)
        else:
            # Tìm trong comment
            for j in i[CommentName]:
                b = j[ValueName].lower()
                if b.find(SearchKey.lower()) != -1:
                    new_results.append(i)
    return new_results


# Hàm phụ trợ: Lấy friends
def GetFriends(uid, page_max, cookie):
    page = 0
    stop = 0
    results = []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    url = 'https://mbasic.facebook.com/search/' + uid + '/friends/'
    while page < page_max and stop == 0:
        page += 1
        result = {}
        r = requests.get(url, headers=headers)
        content = r.content
        # Kiểm tra chặn
        if content.find('Cảnh báo: Vui lòng chậm lại</h3>') != -1:
            # Tìm key
            start1 = content.find('fb_dtsg')
            start2 = content.find('value="', start1) + 7
            end = content.find('"', start2)
            key = content[start2:end]
            content = CrawlFbBypass(url, key, cookie)
        if content.find('Xin lỗi, chúng tôi không thể tìm được bất kỳ kết quả nào cho bạn.</div>') != -1:
            stop = 1
        soup = BeautifulSoup(content, "lxml")
        if len(soup(id="BrowseResultsContainer")) > 0:
            a = str(soup(id="BrowseResultsContainer")[0])
            a = a.split('role="presentation">')
            for i in a[1:]:
                result = {}
                # Lấy tên
                start = i.find('<div class=')
                end = i.find('</div>', start)
                name = i[start:end]
                name = re.sub('<[^>]*>', '', name)
                result['name'] = name
                # Lấy link
                start = i.find('href="') + 6
                end = i.find('__xts__', start)
                link = i[start:end]
                if link.find('&amp;') != -1:
                    link = link.replace('&amp;', '')
                else:
                    link = link.replace('?', '')
                result['link'] = link
                results.append(result)
        # Sang trang mới
        if content.find('"><span>Xem thêm kết quả</span>') != -1:
            end = content.find('"><span>Xem thêm kết quả</span>')
            start = content.rfind('href="', 0, end) + len('href="')
            new_page_link = content[start:end]
            url = new_page_link
        else:
            stop = 1
    return results


# Hàm phụ trợ: Chuyển đổi dữ liệu từ dạng xuất report sang dạng trình diễn
def Report2Present(q):
    results = []
    # Tên [0]
    results.append(['Tên', q['name'], 'name'])
    # Lấy avatar [1]
    results.append(['Đường dẫn ảnh đại diện', q['avatar'], 'avatar'])
    # Lấy cid [2]
    results.append(['cid', q['cid'], 'cid'])
    # Lấy giới thiệu [3]
    results.append(['Giới thiệu', q['about'], 'about'])
    # Lấy giới tính [4]
    results.append(['Giới tính', q['gender'], 'gender'])
    # Lấy biệt danh [5]
    results.append(['Biệt danh', q['nickname'], 'nickname'])
    # Lấy tên khác [6]
    results.append(['Tên khác', q['other_name'], 'other_name'])
    # Lấy ngày sinh [7]
    results.append(['Ngày sinh', q['birthday'], 'birthday'])
    # Lấy ngôn ngữ [8]
    results.append(['Ngôn ngữ', q['language'], 'language'])
    # Lấy nơi sống hiện tại [9]
    results.append(['Nơi sống hiện tại', q['livein'], 'livein'])
    # Lấy quê quán [10]
    results.append(['Quê quán', q['from'], 'from'])
    # Lấy tình trạng quan hệ [11]
    results.append(['Tình trạng quan hệ', q['relationship_status'], 'relationship_status'])
    # Lấy quan điểm tôn giáo [12]
    results.append(['Quan điểm tôn giáo', q['religion'], 'religion'])
    # Lấy quan điểm chính trị [13]
    results.append(['Quan điểm chính trị', q['politics'], 'politics'])
    # Lấy kỹ năng [14]
    results.append(['Kỹ năng', q['skill'], 'skill'])
    # Lấy số điện thoại [15]
    results.append(['Số điện thoại', q['mobile'], 'mobile'])
    # Lấy số điện thoại [16]
    results.append(['Hòm thư điện tử', q['email'], 'email'])
    # Lấy số điện thoại [17]
    results.append(['Trang web', q['website'], 'website'])
    # Lấy các tài khoản trên mạng [18]
    results.append(['Tài khoản trên mạng', q['accounts'], 'accounts'])
    # Lấy các nơi đã làm [19]
    results.append(['Công việc', q['works'], 'works'])
    # Lấy các nơi đã học [20]
    results.append(['Học tập', q['educations'], 'educations'])
    # Lấy các nơi đã học [21]
    results.append(['Pastebin', q['pastebins'], 'pastebins'])
    # Lấy các nơi đã học [22]
    results.append(['Khẩu hiệu', q['tagline'], 'tagline'])
    # Lấy thành tích [23]
    results.append(['Thành tích', q['braggingrights'], 'braggingrights'])
    # Lấy nghề nghiệp [24]
    results.append(['Nghề nghiệp', q['occupation'], 'occupation'])
    # Lấy sự kiện trong đời [25]
    results.append(['Sự kiện trong đời', q['timeline'], 'timeline'])
    # Lấy các dòng trạng thái [26]
    results.append(['Các dòng trạng thái', q['status'], 'status'])
    # Lấy bạn bè [27]
    results.append(['Bạn bè', q['friends'], 'friends'])
    # Lấy sở thích [28]
    results.append(['Sở thích', q['interests'], 'interests'])
    # Lấy các nhóm [29]
    results.append(['Nhóm', q['groups'], 'groups'])
    # Lấy Các trạng thái đối tượng đã bình luận [30]
    results.append(['Các trạng thái đối tượng đã bình luận', q['PublicStatusCommented'], 'PublicStatusCommented'])
    # Lấy nơi đã đến thăm [31]
    results.append(['Các các nơi đã đến thăm', q['LocationVisitedArray'], 'LocationVisitedArray'])
    return results


# /elasticsearch

# Twitter
## Twitter xác thực API
def twitter_api_login():
    global TOKEN_KEY, TOKEN_SECRECT, CONSUMER_KEY, CONSUMER_SECRET
    consumer = oauth2.Consumer(key=CONSUMER_KEY, secret=CONSUMER_SECRET)
    token = oauth2.Token(key=TOKEN_KEY, secret=TOKEN_SECRECT)
    client = oauth2.Client(consumer, token)
    return client


## Twitter API: lấy thông tin từ tên
def twitter_search(q, page, count):
    http_method = "GET"
    post_body = ""
    http_headers = None
    client = twitter_api_login()
    resp, content = client.request(
        'https://api.twitter.com/1.1/users/search.json?q=' + q + '&page=' + str(page) + '&count=' + str(count), \
        method=http_method, body=post_body, headers=http_headers)
    results = json.loads(content)
    return results


## Twitter API: tìm kiếm tài khoản dựa vào username
def twitter_lookup(q):
    http_method = "GET"
    post_body = ""
    http_headers = None
    client = twitter_api_login()
    resp, content = client.request('https://api.twitter.com/1.1/users/lookup.json?screen_name=' + q, \
                                   method=http_method, body=post_body, headers=http_headers)
    results = json.loads(content)
    return results


# Twitter API: Tách kết quả json thành kết quả định dạng của chương trình
def twitter_extract(q):
    ketqua2 = []
    # Check if not exist
    if 'errors' in q:
        print 'return 0'
        return ketqua2
    for i in q:
        ketqua = []
        # name [0]
        ketqua.append(['Tên', i['name'], 'name'])
        # Website [1]
        if 'url' in i['entities']:
            ketqua.append(['Website', i['entities']['url']['urls'][0]['expanded_url'], 'tw_website'])
        else:
            ketqua.append(['Website', '', 'tw_website'])
        # about [2]
        ketqua.append(['Giới thiệu(Twitter)', i['description'], 'tw_about'])
        # location [3]
        ketqua.append(['Địa điểm', i['location'], 'tw_location'])
        # birthday [4] Hien khong duoc API ho tro
        ketqua.append(['Birthday', '', 'tw_birthday'])
        ##flowing,flower number [5][6]
        ketqua.append(['Số người đang theo dõi', i['friends_count'], 'tw_flowing'])
        ketqua.append(['Số người được theo dõi', i['followers_count'], 'tw_flower'])
        # Username [7]
        ketqua.append(['Tài khoản Twitter', i['screen_name'], 'tw_username'])
        # Linkavatar [8]
        ketqua.append(['Liên kết Avatar', i['profile_image_url_https'], 'tw_avatar'])
        # fb_username [9]
        ketqua.append(['Tài khoản Facebook', '', 'fb_username_sub'])
        # Type [10]
        ketqua.append(['Loại thông tin', 'twitter', 'type_information'])
        ketqua2.append(ketqua)
    return ketqua2


# /Twitter

# list FBusername from name by Facebook API Tạm thời bỏ
# def fetch_user_FB(name, offset=0, limit=5000):
#     results = {}
#     ketqua = []
#     token = "EAAAAUaZA8jlABAPsO02oaGiMchD5BQzTzgUSv7XSkFxVQymxb7hL6d9loZCD53iB7j5giI0EUmXSARK1ZCqqSrcakKMxZCxzb2HnPFsaaHpvRMjQNEcBYX0VzW89n29NSS52QfAe1d6fArWgHzhwK8Pxojaqk8DpNio5V9DaUBBinfn2NIQpMyuMnZAOYu3MZD"
#     cookies = []
#     cookie_die = []
#     f = open('cookies_facebook.txt', 'r')
#     for i in f:
#         i = i.strip(' \n\t')
#         i = i.replace(',', ';')
#         cookies.append(i)
#
#     # 01675903806 0963160134 muatrenphobayxa93@gmail.com anhphuongkma khong.triong.1
#     offset = limit * offset  # make offset as page
#     is_correct = 0
#     # dung api de list id facebook can tim
#     r = requests.get('https://graph.facebook.com/v2.11/search?limit=' + str(limit) + '&offset=' + str(
#         offset) + '&type=user&q=' + name + '&access_token=' + token)
#     content = json.loads(r.content)
#     # print content
#     # loop
#     if 'data' not in content:
#         return {'ketqua': '', 'ok': 0}
#     for i in content['data']:
#         is_correct = 1  # set name is correct
#         # Check name
#         # for j in name.split():
#         # 	if i['name'].find(j)==-1:
#         # 		is_correct = 0
#         # 		break
#         if is_correct == 1:
#             ran = randint(0, len(cookies) - 1)
#             cookie = cookies[ran]
#             headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
#                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
#                        'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
#                        'Referer': 'https://m.facebook.com/',
#                        'Origin': 'https://m.facebook.com',
#                        'Cookie': cookie}
#             # lay username, su dung thuat toan random
#             r = requests.head('http://www.facebook.com/' + i['id'], headers=headers)
#             while 'location' not in r.headers:
#                 print 'ohhoho'
#                 ran = randint(0, len(cookies) - 1)
#                 cookie = cookies[ran]
#                 headers = {
#                     'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
#                     'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
#                     'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
#                     'Referer': 'https://m.facebook.com/',
#                     'Origin': 'https://m.facebook.com',
#                     'Cookie': cookie}
#                 r = requests.head('http://www.facebook.com/' + i['id'], headers=headers)
#             username = r.headers['location'].split('.com/')[1]
#             ketqua.append(username)
#     results['ok'] = 1
#     results['ketqua'] = ketqua
#     return results

# Lấy danh sách facebook từ tên, sử dụng Crawler
def fetch_user_FB(q, page_max, cookie):
    page = 0
    stop = 0
    results = []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    url = 'https://mbasic.facebook.com/search/people/?q=' + q
    while page < page_max and stop == 0:
        page += 1
        result = {}
        r = requests.get(url, headers=headers)
        content = r.content
        # Kiểm tra chặn
        if content.find('Cảnh báo: Vui lòng chậm lại</h3>') != -1:
            print 'bi chan'
        if content.find('Rất tiếc, chúng tôi không tìm thấy Trang cho') != -1:
            stop = 1
        soup = BeautifulSoup(content, "lxml")
        if len(soup(id="BrowseResultsContainer")) > 0:
            a = str(soup(id="BrowseResultsContainer")[0])
            a = a.split('role="presentation">')
            for i in a[1:]:
                result = {}
                # Lấy tên
                start = i.find('<div class=')
                end = i.find('</div>', start)
                name = i[start:end]
                name = re.sub('<[^>]*>', '', name)
                result['name'] = name
                # Lấy link
                start = i.find('<a href="') + len('<a href="')
                end = i.find('refid=', start)
                link = i[start:end]
                link = re.sub('<[^>]*>', '', link)
                # Lọc các trường hợp
                link = link.replace('/', '')
                link = link.replace('&amp', '')
                link = link.replace(';', '')
                if link.find('profile.php?id=') == -1:
                    link = link.replace('?', '')
                result['link'] = link
                results.append(result)
        # Sang trang mới
        if content.find('"><span>Xem thêm kết quả</span>') != -1:
            end = content.find('"><span>Xem thêm kết quả</span>')
            start = content.rfind('href="', 0, end) + len('href="')
            new_page_link = content[start:end]
            url = new_page_link
        else:
            stop = 1
    return results


# Fetch FB as requests
def fetch_user_FB_basic(init_infor):
    results = {}
    link_list = []
    ketqua = []
    ketqua1 = []
    global FB_COOKIE
    # 01675903806 0963160134 muatrenphobayxa93@gmail.com anhphuongkma khong.triong.1
    phone = init_infor
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://m.facebook.com/',
               'Origin': 'https://m.facebook.com',
               'Cookie': FB_COOKIE}
    r = requests.get('https://m.facebook.com/search/people/?q=' + phone, headers=headers)
    # If facebook is not be logined
    if r.content.find('href="https://www.facebook.com/login/') != -1:
        results['ok'] = -1
        results['ketqua'] = []
        return results
    for m in re.finditer('\?refid', r.content):
        end_link_pos = m.start()
        start_link_pos = r.content.rfind('/', 0, end_link_pos)
        link = r.content[start_link_pos:end_link_pos]
        link = link.replace('/', '')
        if link.find('/messages/thread/') == -1 and link not in link_list and link != '/' and link != '':
            link_list.append(link)
    return link_list


# FB retrieving function
def FB(q, cookie=FB_COOKIE):
    global FB_COOKIE
    results = {}
    link_list = []
    ketqua = []
    # 01675903806 0963160134 muatrenphobayxa93@gmail.com anhphuongkma khong.triong.1
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://m.facebook.com/',
               'Origin': 'https://m.facebook.com',
               'Cookie': cookie}
    r = requests.get('https://m.facebook.com/search/people/?q=' + q, headers=headers)
    # If facebook is not be logined
    if r.content.find('href="https://www.facebook.com/login/') != -1:
        results['ok'] = -1
        results['ketqua'] = []
        print 'chua login'
        return results

    def removetag(content, tagname):
        while content.find(tagname) != -1:
            start = content.find(tagname)
            end = content.find('>', start) + 1
            content = content.replace(content[start:end], '')
        return content

    # connect vao fb ca nhan

    ketqua = []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://m.facebook.com/search/people/',
               'Origin': 'https://m.facebook.com',
               'Cookie': cookie}
    if q.find('profile.php?') == -1:
        r = requests.get('https://m.facebook.com/' + q + '?v=info', headers=headers)
        r4 = requests.get('https://m.facebook.com/' + q + '?v=likes', headers=headers)
    else:
        r = requests.get('https://m.facebook.com/' + q + '&v=info', headers=headers)
        r4 = requests.get('https://m.facebook.com/' + q + '&v=likes', headers=headers)
    soup = BeautifulSoup(r.content, "lxml")

    # Lấy owner id facebook
    start = r.content.find('owner_id=') + len('owner_id=')
    end = r.content.find('"', start)
    owner_id = r.content[start:end]
    owner_id = owner_id[0:15]
    print 'day la id ' + owner_id

    # lay ten facebook [0]
    if soup.title.text != '':
        a = ['Tên', soup.title.text, 'name']
        ketqua.append(a)

    # lay username facebook	[1]

    a = ['Tài khoản Facebook', q, 'fb_username']
    ketqua.append(a)

    # lay link avatar [2]
    r_avatar = requests.head("https://graph.facebook.com/" + owner_id + "/picture?type=large")
    try:
        link_avatar = r_avatar.headers['location']
    except:
        link_avatar = '/static/images/avatars/no_avatar.jpg'
    # start1 = r.content.find('width="72" height="72"')
    # start2 = r.content.rfind('src="', 0, start1)
    # end = r.content.find('"', start2 + 5)
    # link_avatar = r.content[start2 + 5:end]
    # link_avatar = link_avatar.replace('&amp;', '&')
    a = ['Liên kết hình đại diện', link_avatar, 'fb_avatar']
    ketqua.append(a)

    # lay ten nickname [3]
    start = r.content.find('<span class="alternate_name">')
    if start != -1:
        # Add alternate_name
        end = r.content.find('</span>', start)
        a = r.content[start + len('<span class="alternate_name">'):end]
        alternate_name = a.split('>')[len(a.split('>')) - 1]
        a = ['Biệt danh', alternate_name.decode('utf-8', 'replace')[1:-1], 'nickname']
        ketqua.append(a)
    # Blank alternate_name
    else:
        alternate_name = ''
        a = ['Biệt danh', alternate_name.decode('utf-8', 'replace'), 'nickname']
        ketqua.append(a)

    # lay link facebook [4]
    # start = r.content.find('Facebook</span>')
    # end = r.content.find('</div>',start+30)
    # a =  r.content[start:end]
    # link= a.split('>')[len(a.split('>'))-1]
    # if link!='':
    # 	link = 'https://facebook.com'+link.decode('utf-8', 'replace')
    # 	link = link.replace(' ','')
    # 	a=['link',link,'Link FB']
    # 	ketqua.append(a)
    link = 'https://facebook.com/' + q
    a = ['Liên kết Facebook', link, 'fb_link']
    ketqua.append(a)

    # cong viec [5]
    soup = BeautifulSoup(r.content, "lxml")
    if len(soup(id="work")) > 0:
        work_html = str(soup(id="work")[0])

        work_full = ''
        work_array = []
        for index, val in enumerate(work_html.split('<img ')):
            work_line = ''
            if index != 0:
                a = re.sub('<[^>]*>', ',', val)
                a = re.sub(",+", ", ", a)
                start = a.find('alt')
                end = a.find('"/>', start) + 3
                a = a[end:]

                work_line = a
                work_line = work_line[1:]
            if work_line != '':
                work_full += work_line + ' + '
                # Cho vào mảng
                work_array.append(work_line.split(','))
        work_full = work_full[:-3]  # bo dau phay thua

        a = ['Công việc', work_full.decode('utf-8', 'replace'), 'workin', work_array]
        ketqua.append(a)
    else:
        a = ['Công việc', '', 'workin', []]
        ketqua.append(a)
    # start = r.content.find('Công việc')
    # end1 = r.content.find('<a href=', start)
    # end2 = r.content.find('</a>', end1)
    # a = r.content[start:end2]
    # workat = a.split('>')[len(a.split('>')) - 1]

    # thanh pho hien tai [6]
    start = r.content.find('Th&#xe0;nh ph&#x1ed1; hi&#x1ec7;n t&#x1ea1;i')
    end = r.content.find('</a>', start)
    a = r.content[start:end]
    livein = a.split('>')[len(a.split('>')) - 1]
    a = ['Nơi ở hiện tại', livein.decode('utf-8', 'replace'), 'livein']
    ketqua.append(a)

    # que quan [7]
    start = r.content.find('Qu&#xea; qu&#xe1;n')
    end = r.content.find('</a>', start)
    a = r.content[start:end]
    locationfrom = a.split('>')[len(a.split('>')) - 1]
    a = ['Quê quán', locationfrom.decode('utf-8', 'replace'), 'locationfrom']
    ketqua.append(a)

    # nam sinh [8]
    start = r.content.find('Năm sinh')
    end = r.content.find('</div>', start + 20)
    a = r.content[start:end]
    born = a.split('>')[len(a.split('>')) - 1]
    a = ['Năm sinh', born.decode('utf-8', 'replace'), 'birthday']
    ketqua.append(a)

    # gioi tinh [9]
    start = r.content.find('Giới tính')
    end = r.content.find('</div>', start + 20)
    a = r.content[start:end]
    sex = a.split('>')[len(a.split('>')) - 1]
    a = ['Giới tính', sex.decode('utf-8', 'replace'), 'sex']
    ketqua.append(a)

    # Tinh trang quan he [10]
    start = r.content.find('Tình trạng mối quan hệ')
    end = r.content.find('</div>', start + 100)
    a = r.content[start:end]
    relasionship = a.split('>')[len(a.split('>')) - 1]
    a = ['Tình trạng quan hệ', relasionship.decode('utf-8', 'replace'), 'relasionship']
    ketqua.append(a)

    # Gioi thieu ban than [11]
    start1 = r.content.find('Giới thiệu về ')
    about = ''
    if start1 != -1:
        start2 = r.content.find('<div class=', start1)
        start3 = r.content.find('<div class=', start2)
        start4 = r.content.find('>', start3 + 20) + 1
        end = r.content.find('</div>', start4)
        a = r.content[start4:end]
        a = removetag(a, '<span class=')
        a = removetag(a, '<img class=')
        a = a.replace('</span>', '')
        # Lấy giới thiệu ở đầu trang
        start = r.content.find("""</span></span>""")
        end = r.content.find('<table class=', start)
        NewAbout = r.content[start:end]
        a = a + '. ' + NewAbout
        about = output_html(a)
        a = ['Giới thiệu', about.decode('utf-8', 'replace'), 'about']
        ketqua.append(a)
    else:
        a = ['Giới thiệu', '', 'about']
        ketqua.append(a)
    # ky nang chuyen mon [12]
    start = r.content.find('Kỹ năng chuyên môn')
    end = r.content.find('</span>', start + 110)
    a = r.content[start:end]
    skill = a.split('>')[len(a.split('>')) - 1]
    a = ['Kỹ năng chuyên môn', skill.decode('utf-8', 'replace'), 'skill']
    ketqua.append(a)

    # Hoc van [13]

    soup = BeautifulSoup(r.content, "lxml")
    if len(soup(id="education")) > 0:
        education_html = str(soup(id="education")[0])

        edu_full = ''
        for index, val in enumerate(education_html.split('<img ')):
            edu_line = ''
            edu_array = []
            if index != 0:
                a = re.sub('<[^>]*>', ',', val)
                a = re.sub(",+", ", ", a)
                start = a.find('alt')
                end = a.find('"/>', start) + 3
                a = a[end:]

                edu_line = a
                edu_line = edu_line[1:]

            if edu_line != '':
                edu_full += edu_line + ' + '
                # Cho vào mảng
                edu_array.append(edu_line.split(','))
        edu_full = edu_full[:-3]  # bo dau phay thua

        a = ['Học vấn', edu_full.decode('utf-8', 'replace'), 'education', edu_array]
        ketqua.append(a)
    else:
        a = ['Học vấn', '', 'education', []]
        ketqua.append(a)

    # Biet hieu [14]
    start = r.content.find('Biệt hiệu')
    end = r.content.find('</div>', start + 30)
    a = r.content[start:end]
    nickname = a.split('>')[len(a.split('>')) - 1]
    a = ['Biệt hiệu', nickname.decode('utf-8', 'replace'), 'nickname']
    ketqua.append(a)

    # SDT [15]
    start = r.content.find('Di động')
    end = r.content.find('</span>', start + 30)
    a = r.content[start:end]
    phone = a.split('>')[len(a.split('>')) - 1]
    a = ['Số điện thoại', phone.decode('utf-8', 'replace'), 'phone']
    ketqua.append(a)

    # ngay sinh [16]
    start = r.content.find('Ngày sinh')
    end = r.content.find('</div>', start + 30)
    a = r.content[start:end]
    birthday = a.split('>')[len(a.split('>')) - 1]
    a = ['Ngày sinh', birthday.decode('utf-8', 'replace'), 'birthday']
    ketqua.append(a)

    # ngon ngu [17]
    start = r.content.find('Ngôn ngữ')
    end = r.content.find('</div>', start + 30)
    a = r.content[start:end]
    language = a.split('>')[len(a.split('>')) - 1]
    a = ['Ngôn ngữ', language.decode('utf-8', 'replace'), 'language']
    ketqua.append(a)

    # yahoo [18]
    start = r.content.find('Yahoo! Messenger</span>')
    end = r.content.find('</div>', start + 30)
    a = r.content[start:end]
    yahoo = a.split('>')[len(a.split('>')) - 1]
    a = ['Tài khoản Yahoo!', yahoo.decode('utf-8', 'replace'), 'yahoo']
    ketqua.append(a)

    # Quan diem ton giao [19]
    start = r.content.find('Quan điểm tôn giáo')
    end = r.content.find('</a>', start)
    a = r.content[start:end]
    religion = a.split('>')[len(a.split('>')) - 1]
    a = ['Quan điểm tôn giáo', religion.decode('utf-8', 'replace'), 'religion']
    ketqua.append(a)

    # Quan diem chinh tri [20]
    if r.content.find('Quan Điểm Chính Trị') != -1:
        start = r.content.find('Quan Điểm Chính Trị') + len('Quan Điểm Chính Trị')
        end = r.content.find('</div></td></tr>', start)
        a = r.content[start:end]
        politics = re.sub('<[^>]*>', '', a)
        a = ['Quan điểm chính trị', politics.decode('utf-8', 'replace'), 'politics']
        ketqua.append(a)
    else:
        a = ['Quan điểm chính trị', '', 'politics']
        ketqua.append(a)

    # Email [21]
    start = r.content.find('Email</span>')
    end = r.content.find('</a>', start)
    a = r.content[start:end]
    email = a.split('>')[len(a.split('>')) - 1]
    a = ['Hòm thư điện tử', email.decode('utf-8', 'replace'), 'email']
    ketqua.append(a)

    # Trang web [22]
    start = r.content.find('Trang web</span>')
    end = r.content.find('</a>', start)
    a = r.content[start:end]
    website = a.split('>')[len(a.split('>')) - 1]
    a = ['Trang web', website.decode('utf-8', 'replace'), 'website']
    ketqua.append(a)

    # Twitter username [23]
    a = ['Tài khoản Twitter', '', 'tw_username_sub', []]
    ketqua.append(a)

    # Địa chỉ [24]
    start = r.content.find('Địa chỉ</span>')
    end = r.content.find('</a>', start)
    a = r.content[start:end]
    address = a.split('>')[len(a.split('>')) - 1]
    a = ['Địa chỉ', address.decode('utf-8', 'replace'), 'address', []]
    ketqua.append(a)
    print 'da xong 1'
    # Group công khai, đưa ra link [25]
    start = r.content.find('?owner_id=')
    end = r.content.find('"', start)
    id_fb = r.content[start + 10:end]
    public_group = 'https://www.facebook.com/search/' + id_fb + '/groups'
    a = ['Nhóm công khai (FB)', public_group.decode('utf-8', 'replace'), 'public_group']
    ketqua.append(a)
    print 'da xong 2'
    # Sự kiện trong đời [26]
    start = r.content.find('Sự kiện trong đời')
    new_content = r.content[start + 24:]
    timeline_array = []
    if new_content.find('width="16" height="16"') != -1:
        end = new_content.find('width="16" height="16"')
        while new_content.find('width="16" height="16"', end + 1) != -1:
            end = new_content.find('width="16" height="16"', end + 1)
        end = new_content.find('</div>', end)
        new_content = new_content[:end]
        timeline = re.sub('<[^>]*>', ',', new_content)
        timeline_for_array = re.sub(",+", "!|!", timeline)  # Sử dụng cho timeline array
        timeline = re.sub(",+", ", ", timeline)
        timeline = timeline[2:]  # Bỏ dấu đầu tiên đi

        # timeline array
        year = 0
        timeline_for_array = timeline_for_array[3:]  # Bỏ dấu đầu tiên đi
        for i in timeline_for_array.split('!|!'):
            if i.isdigit() and len(i) == 4:
                year = i
            elif i.replace(' ', '') != '':
                timeline_array.append({'year': year, 'value': i})

        a = ['Sự kiện trong đời', timeline.decode('utf-8', 'replace'), 'timeline', timeline_array]
        ketqua.append(a)
    else:
        a = ['Sự kiện trong đời', '', 'timeline', timeline_array]
        ketqua.append(a)
    print 'da xong 3'
    # Xem add friend chưa [27]
    if r.content.find('/a/mobile/friends/profile_add_friend.php') == -1:
        a = ['Đã add friend chưa?', 'ok', 'is_fb_friend']
        ketqua.append(a)
    else:
        a = ['Đã add friend chưa?', 'no', 'is_fb_friend']
        ketqua.append(a)
    print 'da xong 4'
    # Lấy status facebook [28]
    fb_status_array = GetFbStatus(q, cookie, 1)['results']
    a = ['Trạng thái facebook', fb_status_array, 'fb_status', fb_status_array]
    ketqua.append(a)

    # loại thông tin [29]
    a = ['Loại thông tin', 'facebook', 'type_information']
    ketqua.append(a)

    # Friends đại diện [30]
    friend_array = GetFriends(owner_id, 2, cookie)
    a = ['Bạn bè', 'Xem trong báo cáo', 'friends', friend_array]
    ketqua.append(a)
    print 'da xong 5'
    # Sở thích [31]
    interests_array = []
    for i in r4.content.split('<h4 class=')[1:]:
        i = '<' + i  # Đóng thẻ cho thẻ h4 đã split
        end = i.find('</h4>')
        cate = i[0:end]
        cate = re.sub('<[^>]*>', '!|!', cate)
        cate = cate.replace('!|!', '')
        start = i.find('<span>')
        while start != -1:
            end = i.find('</span>', start)
            name = i[start:end]
            name = re.sub('<[^>]*>', '!|!', name)
            name = name.replace('!|!', '')
            if name != 'Xem thêm':
                interests_array.append({'cate': cate, 'name': name})
            start = i.find('<span>', end)
    # Cho hết vào một category
    new_interests_array = []
    cate = []
    for i in interests_array:
        if i['cate'] not in cate:
            cate.append(i['cate'])
    print cate
    for i in cate:
        name = ""
        for j in interests_array:
            if j['cate'] == i:
                name += j['name'] + ', '
        name = name[:-2]
        new_interests_array.append({'cate': i, 'name': name})
    a = ['Sở thích', 'Xem trong báo cáo', 'interests', new_interests_array]
    ketqua.append(a)
    print 'da xong 6'
    # Lấy hết Group [32]
    groups_array = GetFbGroup(owner_id, cookie)
    a = ['Nhóm', groups_array, 'groups', groups_array]
    ketqua.append(a)
    print 'da xong 7'
    # Lấy các post mà đối tượng comment [33]
    commentpost_array = GetCommentPost(owner_id, 10, cookie)
    a = ['Các bài viết đã bình luận', commentpost_array, 'commentpost', commentpost_array]
    ketqua.append(a)
    print 'da xong 8'
    # Lấy các địa điểm đã đến thăm [34]
    LocationVisitedArray = GetPlaceVisited(owner_id, 10, cookie)
    a = ['Các bài viết đã bình luận', LocationVisitedArray, 'LocationVisitedArray', LocationVisitedArray]
    ketqua.append(a)
    print 'da xong 9'
    results['ok'] = 1
    results['ketqua'] = ketqua
    return results


def _facebook_to_compare(profile):
    results = {}
    if 'ketqua' in profile:
        profile = profile['ketqua']
    results['about'] = profile[11][1]
    results['website'] = profile[22][1]
    results['email'] = profile[21][1]
    results['phone'] = profile[15][1]
    results['website_about'] = detect_website_link(results['about'])
    results['email_about'] = ''  # chua co thuat toan
    results['location'] = no_accent_vietnamese(profile[6][1].encode('utf-8')).lower()
    results['from'] = no_accent_vietnamese(profile[7][1].encode('utf-8')).lower()
    results['name'] = profile[0][1]
    results['nickname'] = profile[3][1]
    results['username'] = profile[1][1].lower()
    return results


def _twitter_to_compare(profile):
    results = {}
    results['about'] = profile[2][1]
    results['website'] = profile[1][1]
    results['website_about'] = detect_website_link(results['about'])
    results['email_about'] = ''  # chua co thuat toan
    results['location'] = no_accent_vietnamese(profile[3][1].encode('utf-8')).lower()
    results['name'] = profile[0][1]
    results['username'] = profile[7][1].lower()
    return results


def _generrate_username_list(short_profile):
    username_list = []

    username_list.append(short_profile['username'])  ###Load username

    ###username remove '.'
    a = short_profile['username'].replace('.', '')
    if a not in username_list:
        username_list.append(a)

    ###remove '.'s and numbers on the last of string
    a = re.sub(r'\d+$', '', short_profile['username'].replace('.', ''))
    if a not in username_list:
        username_list.append(a)

    ### No accent vietnamese and remove spaces for the name
    a = no_accent_vietnamese(short_profile['name'].encode('utf-8'))
    a = a.replace(' ', '')
    a = a.lower()
    if a not in username_list:
        username_list.append(a)

    return username_list


def _compare_make_score(short_profile1, short_profile2):
    score = 0
    # an luon
    ## website
    if (short_profile1['website'] == short_profile2['website'] and short_profile1['website'] != '') \
            or (short_profile1['website'] == short_profile2['website_about'] and short_profile2['website_about'] != '') \
            or (short_profile1['website_about'] == short_profile2['website'] and short_profile1['website_about'] != '') \
            or (short_profile1['website_about'] == short_profile2['website_about'] and short_profile1[
        'website_about'] != ''):
        print 'Trung website'
        return 100

    ## email
    if 'email' in short_profile1:
        if (short_profile1['email'] == short_profile2['email_about'] and short_profile1['email'] != '') \
                or (
                short_profile1['email_about'] == short_profile2['email_about'] and short_profile1['email_about'] != ''):
            print 'Trung email'
            return 100
    else:
        if (short_profile2['email'] == short_profile1['email_about'] and short_profile2['email'] != '') \
                or (
                short_profile1['email_about'] == short_profile2['email_about'] and short_profile1['email_about'] != ''):
            print 'Trung email'
            return 100

    # Tinh diem
    ## location
    if 'from' in short_profile1:
        if (short_profile1['from'] == short_profile2['location'] and short_profile1['from'] != '') \
                or (short_profile1['location'] == short_profile2['location'] and short_profile1['location'] != ''):
            print 'Trung location'
            score += 1
    else:
        if (short_profile2['from'] == short_profile1['location'] and short_profile2['from'] != '') \
                or (short_profile2['location'] == short_profile1['location'] and short_profile2['location'] != ''):
            print 'Trung localtion'
            score += 1

    ## name
    if 'nickname' in short_profile1:
        if (short_profile1['nickname'] == short_profile2['name'] and short_profile1['nickname'] != '') \
                or (short_profile1['name'] == short_profile2['name'] and short_profile1['name'] != ''):
            print 'Trung name'
            score += 3
    else:
        if (short_profile2['nickname'] == short_profile1['name'] and short_profile2['nickname'] != '') \
                or (short_profile2['name'] == short_profile1['name'] and short_profile2['name'] != ''):
            print 'Trung name'
            score += 3
    ## username
    if (short_profile1['username'] == short_profile2['username'] and short_profile1['username'] != ''):
        print 'Trung username'
        score += 4

    # Tinh %
    percent = (score * 100) / 10
    return int(percent)


# Find twitter from FB
def fb_to_tw(profile):
    username_list2 = ''
    results = []
    fb_compare = _facebook_to_compare(profile)
    username_list = _generrate_username_list(fb_compare)
    print username_list
    # username cach nhau dau phay
    for _u in username_list:
        username_list2 += _u + ','
    username_list2 = username_list2[:-1]  # bo dau phay thua
    # Look up
    tw_results = twitter_lookup(username_list2)
    tw_results = twitter_extract(tw_results)

    if tw_results == []:
        return profile
    for i in tw_results:
        print i
        tw_compare = _twitter_to_compare(i)
        compare_make_score = _compare_make_score(tw_compare, fb_compare)
        if compare_make_score > 0:
            results.append([tw_compare['username'], compare_make_score, i])
    # tinh max
    if len(results) > 0:
        profile.append(['Ratio', results, 'ratio'])

    return profile


# Find FB from twitter
def tw_to_fb(profile):
    username_list = []
    results = []
    tw_compare = _twitter_to_compare(profile)
    print username_list
    # Truong hop dac biet
    # #If twitter's website is facebook link, it will be his facebook
    if tw_compare['website'] != '':
        print tw_compare['website']
        o = urlparse(tw_compare['website'])
        print o.hostname
        print o.path
        if str(o.hostname) in ['facebook.com', 'fb.com', 'www.facebook.com']:
            path = o.path
            path = path.replace('/', '')
            # Lay thong tin de sau con merge
            fb_result = FB(path)
            # Luu vao result
            results.append([path, 100, fb_result])

    username_list.append(tw_compare['username'])
    # Use username_list
    for _u in username_list:
        fb_result = FB(_u)
        if fb_result['ketqua'][0][1] != 'Không tìm thấy trang':
            fb_compare = _facebook_to_compare(fb_result)
            compare_make_score = _compare_make_score(tw_compare, fb_compare)
            if compare_make_score > 0:
                results.append([tw_compare['username'], compare_make_score, fb_result])
    # tinh max
    if len(results) > 0:
        # sua profile dua ke qua tong quan
        profile.append(['Ratio', results, 'ratio'])
    print results
    return profile


# Xử lý ảnh cho các loại báo cáo
def image_for_report(image_link):
    image_from_url = urllib2.urlopen(image_link)
    print image_from_url.getcode()
    io_url = StringIO.StringIO()
    io_url.write(image_from_url.read())
    io_url.seek(0)
    return io_url


# Xử lý thông tin cho báo cáo tiêu chuẩn
def data_for_stdreport(profile, user_system, tpl, WebMode=0):
    ## Khai báo
    results = {}
    results['name'] = ''
    results['cid'] = ''
    results['ckid'] = ''
    results['about'] = ''
    results['gender'] = ''
    results['nickname'] = ''
    results['other_name'] = ''
    results['birthday'] = ''
    results['language'] = ''
    results['livein'] = ''
    results['from'] = ''
    results['relationship_status'] = ''
    results['religion'] = ''
    results['politics'] = ''
    results['skill'] = ''
    results['mobile'] = ''
    results['email'] = ''
    results['website'] = ''
    results['accounts'] = []
    results['works'] = []
    results['educations'] = []
    results['pastebins'] = []
    results['tagline'] = ''
    results['braggingrights'] = ''
    results['occupation'] = ''
    results['avatar'] = ''
    results['timeline'] = []
    results['status'] = []
    results['friends'] = []
    results['interests'] = []
    results['groups'] = []
    results['PublicStatusCommented'] = []
    results['LocationVisitedArray'] = []
    io_url = ''
    results['LocalAccounts'] = []
    ## Chắt lọc thông tin từ các profile
    for i in profile:
        ### Facebook
        if find_data(i, 'type_information')[1] == 'facebook':
            # Tên lấy một thôi
            results['name'] = advance_plus_string(results['name'], i[0][1], 1)
            results['about'] = advance_plus_string(results['about'], i[11][1], 0, '. ')
            # Giới tính lấy một thôi
            results['gender'] = advance_plus_string(results['gender'], i[9][1], 1)
            results['nickname'] = advance_plus_string(results['nickname'], i[3][1])
            results['other_name'] = advance_plus_string(results['other_name'], i[0][1])
            results['birthday'] = advance_plus_string(results['birthday'], i[16][1])
            results['language'] = advance_plus_string(results['language'], i[17][1])
            results['livein'] = advance_plus_string(results['livein'], i[6][1], 0, '. ')
            results['from'] += i[7][1]
            results['relationship_status'] = advance_plus_string(results['relationship_status'], i[10][1], 1)
            results['religion'] += i[19][1]
            results['politics'] += i[20][1]
            results['skill'] = advance_plus_string(results['skill'], i[12][1])
            results['mobile'] = advance_plus_string(results['mobile'], i[15][1])
            results['email'] = advance_plus_string(results['email'], i[21][1])
            results['website'] = advance_plus_string(results['website'], i[22][1])
            results['works'] += i[5][3]
            results['educations'] += i[13][3]
            results['timeline'] += i[26][3]
            results['status'] += i[28][3]
            results['friends'] += i[30][3]
            results['interests'] += i[31][3]
            results['accounts'].append({'value': 'http://facebook.com/' + i[1][1], 'type': 'Facebook'})
            results['groups'] += i[32][3]
            results['PublicStatusCommented'] += i[33][3]
            results['LocationVisitedArray'] += i[34][3]
            ## Xử lý avatar, lấy cái đầu tiên thôi
            if WebMode == 0:
                if io_url == '':
                    io_url = image_for_report(i[2][1])
                    myimage = InlineImage(tpl, io_url, width=Inches(1.5))
                    results['avatar'] = myimage
            else:
                results['avatar'] = i[2][1]

        ### Twitter
        if find_data(i, 'type_information')[1] == 'twitter':
            results['name'] = advance_plus_string(results['name'], i[0][1], 1)
            results['other_name'] = advance_plus_string(results['other_name'], i[0][1])
            results['accounts'].append({'value': 'http://twitter.com/' + i[7][1], 'type': 'Twitter'})
            results['about'] = advance_plus_string(results['about'], i[2][1], 0, '. ')
            results['website'] = advance_plus_string(results['website'], i[1][1])
            results['livein'] = advance_plus_string(results['livein'], i[3][1], 0, '. ')
            if WebMode == 0:
                if io_url == '':
                    io_url = image_for_report(i[8][1])
                    myimage = InlineImage(tpl, io_url, width=Inches(1.5))
                    results['avatar'] = myimage
            else:
                results['avatar'] = i[8][1]

        ### Google plus
        if find_data(i, 'type_information')[1] == 'googleplus':
            results['name'] = advance_plus_string(results['name'], i[0][1], 1)
            results['other_name'] = advance_plus_string(results['other_name'], i[0][1])
            results['about'] = advance_plus_string(results['about'], i[4][1], 0, '. ')
            results['skill'] = advance_plus_string(results['skill'], i[10][1])
            results['livein'] = advance_plus_string(results['livein'], i[9][1], 0, '. ')
            results['nickname'] = advance_plus_string(results['nickname'], i[12][1])
            results['birthday'] = advance_plus_string(results['birthday'], i[11][1], 1)
            results['tagline'] = advance_plus_string(results['tagline'], i[6][1], 0, '. ')
            results['occupation'] = advance_plus_string(results['occupation'], i[3][1])
            results['braggingrights'] = advance_plus_string(results['braggingrights'], i[5][1], 0, '. ')
            results['works'].append([i[8][1]])
            results['accounts'].append({'value': 'http://plus.google.com/' + i[1][1], 'type': 'Google Plus'})
            if WebMode == 0:
                if io_url == '':
                    io_url = image_for_report(i[2][1])
                    myimage = InlineImage(tpl, io_url, width=Inches(1.5))
                    results['avatar'] = myimage
            else:
                results['avatar'] = i[2][1]

        ### Blogger
        if find_data(i, 'type_information')[1] == 'blogger':
            results['name'] = advance_plus_string(results['name'], i[1][1], 1)
            results['other_name'] = advance_plus_string(results['other_name'], i[1][1])
            results['about'] = advance_plus_string(results['about'], i[8][1], 0, '. ')
            results['livein'] = advance_plus_string(results['livein'], i[3][1], 0, '. ')
            results['email'] = advance_plus_string(results['email'], i[2][1])
            results['occupation'] = advance_plus_string(results['occupation'], i[9][1])
            results['website'] = advance_plus_string(results['website'], i[10][1])
            results['accounts'].append({'value': 'http://blogger.com/' + i[12][1], 'type': 'blogger'})
            if WebMode == 0:
                if io_url == '':
                    io_url = image_for_report(i[0][1])
                    myimage = InlineImage(tpl, io_url, width=Inches(1.5))
                    results['avatar'] = myimage
            else:
                results['avatar'] = i[0][1]

        ### Pastebin
        if find_data(i, 'type_information')[1] == 'pastebin':
            results['pastebins'].append(['https://pastebin.com' + i[3][1], i[1][1]])

        print i
        ### Dữ liệu nội bộ
        if find_data(i, 'type_information')[1] == 'cred_local_data':
            results['LocalAccounts'].append({'title': i[2][1], 'username': i[0][1], 'password': i[1][1]})

    ## Bổ sung các thông tin khác
    results['user_system'] = user_system
    for j in i:
        if j[2] == 'cid':
            results['cid'] = j[1]
        if j[2] == 'ckid':
            results['ckid'] = j[1]
    ### Làm giới thiệu tóm tắt
    brief = ""
    if results['name'] != '':
        brief += u"Tên tôi là " + results['name'] + ". "
    if results['gender'] != '':
        brief += u"Tôi là " + results['gender'] + ". "
    # if results['livein'] != '' and results['from'] == '':
    #     brief += u"Tôi đến từ " + results['livein'] + ". "
    # if results['livein'] == '' and results['from'] != '':
    #     brief += u"Tôi đến từ " + results['from'] + ". "
    if len(results['educations']) > 0:
        brief += u"Tôi học tại " + results['educations'][0][0] + " "
    if len(results['works']) > 0:
        if len(results['educations']) == 0:
            brief += u"Tôi làm tại " + results['works'][0][0]
        else:
            brief += u"và tôi làm tại " + results['works'][0][0]
    ### Tài khoản trên mạng
    account_brief = ""
    if len(results['accounts']) > 0:
        account_brief += u"Tôi đã đăng ký các tài khoản trên "
        for i in results['accounts']:
            if account_brief.find(i['type']) == -1:
                account_brief += i['type'] + ', '
        account_brief = account_brief[:-2] + "."
    results['brief'] = brief
    results['account_brief'] = account_brief
    return results


# Mẫu báo cáo tiêu chuẩn
def report_stardard(profile, filename, user_system):
    ## Nạp mẫu
    tpl = DocxTemplate('template.docx')

    context = data_for_stdreport(profile, user_system, tpl)

    tpl.render(context)
    tpl.save('downloads/' + filename + '.docx')
    result = filename + '.docx'
    return result


# Mẫu báo cáo kỹ thuật
def report_tech(profile, filename):
    document = Document()
    image_from_url = urllib2.urlopen(profile[2][1])
    io_url = StringIO.StringIO()
    io_url.write(image_from_url.read())
    io_url.seek(0)
    try:
        document.add_picture(io_url, width=Inches(1.25))
    except:
        pass
    for i in profile:
        i[1] = str(i[1])
        p = document.add_paragraph(unicode(i[0].encode("utf-8") + ':' + i[1].encode("utf-8"), 'utf-8'))
    document.save('downloads/' + filename + '.docx')
    result = filename + '.docx'
    return result


# Download reports
def download(profile, type, user_system):
    link = ""
    # generate name
    uuid_ran = str(uuid.uuid1())
    hash_ran = str(random.getrandbits(128))
    final_ran = hashlib.sha224(uuid_ran + hash_ran).hexdigest()
    # save file
    if type == 'std':
        link = report_stardard(profile, final_ran, user_system)
    elif type == 'tech':
        link = report_tech(profile, final_ran)
    return link


# Gộp thông tin hai tài khoản
def merg_fb_tw(fb_profile, tw_profile):
    print fb_profile
    print tw_profile
    # Ghép username fb
    if tw_profile[9][1] != '' and fb_profile[1][1] == '':
        fb_profile[1][1] += ', ' + tw_profile[9][1]
        fb_profile[1][1] = fb_profile[1][1].strip(' ,')
    if tw_profile[23][1] != '' and fb_profile[23][1] == '':
        # Ghép username twitter
        fb_profile[23][1] += ', ' + tw_profile[23][1]
        fb_profile[23][1] = fb_profile[23][1].strip(' ,')
    # Ghép địa chỉ website
    if tw_profile[1][1] != '':
        fb_profile[22][1] += ', ' + tw_profile[1][1]
        fb_profile[22][1] = fb_profile[22][1].strip(' ,')
    # Ghép Địa điểm
    if tw_profile[3][1] != '':
        fb_profile[6][1] += ', ' + tw_profile[3][1]
        fb_profile[6][1] = fb_profile[6][1].strip(' ,')
    # Ghép giới thiệu
    if tw_profile[8][1] != '':
        fb_profile[11][1] += ', ' + tw_profile[8][1]
        fb_profile[11][1] = fb_profile[11][1].strip(' ,')
        fb_profile[11][1] = fb_profile[11][1].replace('&#039;', '')
    return fb_profile


# Viết lại các dữ liệu cho các dữ liệu mảng Google plus
def gp_dict_data_handle(q):
    result = ""
    results = ""
    for i in q:
        result = ""
        if 'title' in i:
            result += i['title']
            if 'name' in i:
                result += u" tại " + i['name']
                if 'startDate' in i:
                    result += u" bắt đầu từ " + i['startDate']
                    if 'endDate' in i:
                        result += u" đến " + i['endDate']
            results += result + u", "
        elif 'name' in i:
            result = ""
            result += i['name']
            results += result + u", "
        if 'value' in i:
            result = ""
            result = i['value']
            results += result + u", "
    return results[:-2]


# Google Plus
def google_plus(q):
    global GGPL_KEY
    results = {}
    ketqua = []
    r = requests.get("https://www.googleapis.com/plus/v1/people/" + q + "?key=" + GGPL_KEY)
    time.sleep(2)
    content = json.loads(r.content)
    if 'objectType' in content:
        if content['objectType'] != 'person':
            results['ok'] = 0
            results['ketqua'] = []
            return results
    else:
        results['ok'] = 0
        results['ketqua'] = []
        return results
    # Name [0]
    a = ['Tên', content['displayName'].encode('utf-8', 'replace').decode('utf-8'), 'name']
    ketqua.append(a)

    # Username [1]
    a = ['Tài khoản Google Plus', q, 'gp_username']
    ketqua.append(a)

    # Avatar [2]
    a = ['Liên kết hình đại diện', content['image']['url'].decode('utf-8', 'replace'), 'gp_avatar']
    ketqua.append(a)

    # Nghề nghiệp [3]
    if 'occupation' in content:
        a = ['Nghề nghiệp', content['occupation'].encode('utf-8', 'replace').decode('utf-8'), 'occupation']
    else:
        a = ['Nghề nghiệp', '', 'occupation']
    ketqua.append(a)

    # Giới thiệu về bản thân tôi [4]
    if 'aboutMe' in content:
        content['aboutMe'] = output_html(content['aboutMe'])
        a = ['Giới thiệu', content['aboutMe'], 'about']
    else:
        a = ['Giới thiệu', '', 'occupation']
    ketqua.append(a)

    # Thành tích  [5]
    if 'braggingRights' in content:
        a = ['Thành tích', content['braggingRights'], 'braggingright']
    else:
        a = ['Giới thiệu', '', 'occupation']
    ketqua.append(a)

    # Khẩu hiệu [6]
    if 'tagline' in content:
        a = ['Khẩu hiệu', content['tagline'].encode('utf-8', 'replace').decode('utf-8'), 'tagline']
    else:
        a = ['Khẩu hiệu', '', 'tagline']
    ketqua.append(a)

    # Relasionship [7]
    if 'relationshipStatus' in content:
        a = ['Tình trạng quan hệ', content['relationshipStatus'], 'relasionship']
    else:
        a = ['Tình trạng quan hệ', '', 'relasionship']
    ketqua.append(a)

    # Những nơi đã làm việc [8]
    if 'organizations' in content:
        content['organizations'] = gp_dict_data_handle(content['organizations'])
        a = ['Nơi làm việc', content['organizations'], 'workin']
    else:
        a = ['Nơi làm việc', '', 'workin']
    ketqua.append(a)

    # Những nơi đã ở [9]
    if 'placesLived' in content:
        content['placesLived'] = gp_dict_data_handle(content['placesLived'])
        a = ['Nơi ở hiện tại', content['placesLived'], 'livein']
    else:
        a = ['Nơi ở hiện tại', '', 'livein']
    ketqua.append(a)

    # Kỹ năng [10]
    if 'skills' in content:
        a = ['Kỹ năng', content['skills'], 'skill']
    else:
        a = ['Kỹ năng', '', 'skill']
    ketqua.append(a)

    # Ngày sinh [11]
    if 'birthday' in content:
        a = ['Ngày sinh', content['birthday'], 'birthday']
    else:
        a = ['Ngày sinh', '', 'birthday']
    ketqua.append(a)

    # Nickname [12]
    if 'nickname' in content:
        a = ['Biệt danh', content['nickname'].encode('utf-8', 'replace').decode('utf-8'), 'nickname']
    else:
        a = ['Biệt danh', '', 'nickname']
    ketqua.append(a)

    # Loại thông tin [13]
    a = ['Loại thông tin', 'googleplus', 'type_information']
    ketqua.append(a)

    results['ok'] = 1
    results['ketqua'] = ketqua
    return results


# Google plus fetch
def google_plus_fetch(q, first):
    results = []
    ketqua = bing_web_related(q, 'plus.google.com', first)['ketqua']
    for i in ketqua:
        u = urlparse(i)
        b = u.path.replace('/', '')
        a = google_plus(b)
        if a['ok'] == 1:
            results.append(a['ketqua'])
    return results


# All webs related
def all_web_related(q):
    results = {}
    all_web_results = []
    page_number = 1  # Tạm thời để 1 page
    print page_number
    # try:
    r = requests.get('http://www.bing.com/search?q=site:plus.google.com ' + q + '&first=' + str(page_number))
    soup = BeautifulSoup(r.content, "lxml")
    mydivs = soup.findAll("li", {"class": "b_algo"})
    for mydiv in mydivs:
        # print mydiv.a.attrs['href']
        # r1 = requests.get(mydiv.a.attrs['href'])
        # r1.content.find(q.decode('utf-8', 'replace')) != -1 and
        if mydiv.a.attrs['href'].decode('utf-8', 'replace') not in all_web_results:
            gp_web_results.append(mydiv.a.attrs['href'])
    # except:
    #     print 'Connection Error!'
    a = {'ok': 1, 'ketqua': all_web_results}
    results.update(a)
    return results


# Hỗ trợ lấy nội dung từ blogger
def blogger_fetch_html(content, patterm, end_patterm='</span>'):
    if content.find(patterm) != -1:
        start = content.find(patterm) + len(patterm)
        end = content.find(end_patterm, start)
        ok = content[start:end]
        ok = re.sub('<[^>]*>', '', ok)
    else:
        ok = ""
    return ok


# Hàm chính blogger
def blogger(q):
    ketqua = []
    r = requests.get('https://www.blogger.com/' + q)

    # Lấy avatar [0]
    soup = BeautifulSoup(r.content, "lxml")
    ava = str(soup.find_all("img", class_="photo", attrs="src")[0]['src'])
    if ava == "/img/avatar_blue_m_96.png":
        ava = "//www.blogger.com/img/avatar_blue_m_96.png"
    a = ['Liên kết hình đại diện', ava, 'blg_avatar']
    ketqua.append(a)

    # Lấy tên [1]
    name = str(soup.h1.text.encode('utf-8', 'replace'))
    a = ['Tên', name, 'name']
    ketqua.append(a)

    # Lấy email [2]
    if r.content.find('printEmail(') != -1:
        start = r.content.find('printEmail(') + 12
        end = r.content.find('"', start)
        email = r.content[start:end]
    else:
        email = ""
    a = ['Hòm thư điện tử', email, 'email']
    ketqua.append(a)

    # Lấy địa điểm [3]
    livein = blogger_fetch_html(r.content, '<tr><th class="item-key">Vị trí</th>', "</span></td></tr>")
    livein = stringToConvert(livein)
    a = ['Địa điểm', livein, 'livein']
    ketqua.append(a)

    # Lấy sở thích [4]
    interesting = blogger_fetch_html(r.content, '<tr><th class="item-key">Sở thích</th>')
    interesting = stringToConvert(interesting)
    a = ['Sở thích', interesting, 'interesting']
    ketqua.append(a)

    # Lấy phim yêu thích [5]
    film_favor = blogger_fetch_html(r.content, '<tr><th class="item-key">Phim Ưa thích</th>')
    film_favor = stringToConvert(film_favor)
    a = ['Phim yêu thích', film_favor, 'film_favor']
    ketqua.append(a)

    # Lấy nhạc yêu thích [6]
    music_favor = blogger_fetch_html(r.content, '<tr><th class="item-key">Nhạc Ưa thích</th>')
    music_favor = stringToConvert(music_favor)
    a = ['Nhạc yêu thích', music_favor, 'music_favor']
    ketqua.append(a)

    # Lấy sách yêu thích [7]
    book_favor = blogger_fetch_html(r.content, '<tr><th class="item-key">Sách Ưa thích</th>')
    book_favor = stringToConvert(book_favor)
    a = ['Sách yêu thích', book_favor, 'book_favor']
    ketqua.append(a)

    # Lấy giới thiệu [8]
    about = blogger_fetch_html(r.content, '<tr><th class="item-key">Giới thiệu</th>', '</td></tr>')
    about = stringToConvert(about)
    a = ['Giới thiệu', about, 'about']
    ketqua.append(a)

    # Lấy nghề nghiệp [9]
    job = blogger_fetch_html(r.content, '<tr><th class="item-key">Nghề nghiệp</th>')
    job = stringToConvert(job)
    a = ['Nghề nghiệp', job, 'job']
    ketqua.append(a)

    # Lấy trang web [10]
    if r.content.find('me nofollow') != -1:
        end1 = r.content.find('me nofollow')
        start = r.content.rfind('href="', 0, r.content.find('me nofollow')) + 6
        end2 = r.content.find('"', start)
        website = r.content[start:end2]
    else:
        website = ""
    a = ['Trang web', website, 'website']
    ketqua.append(a)

    # Loại thông tin [11]
    a = ['Loại thông tin', 'blogger', 'type_information']
    ketqua.append(a)

    # URI (Đường dẫn URI) [12]
    a = ['Đường dẫn URI', q, 'blogger_uri']
    ketqua.append(a)

    # Return
    return ketqua


# Blogger fetch
def blogger_fetch(q, first):
    results = []
    ketqua = bing_web_related(q, 'blogger.com instreamset:url:profile', first)['ketqua']
    for i in ketqua:
        print i
        u = urlparse(i)
        b = u.path
        print b
        # Kiểm tra xem link có profile không
        if b.find('profile') != -1:
            a = blogger(b)
            results.append(a)
    return results


# Hàm chính extract pastebin
def pastebin(b, q):
    ketqua = []
    r = requests.get('http://pastebin.com/raw' + b)

    # Link [0]
    a = ['Đường dẫn', 'http://pastebin.com/raw' + b, 'pastebin_url']
    ketqua.append(a)

    # Nội dung vắn tắt [1]
    rlow = r.content.lower()
    qlow = q.lower()
    vitri = rlow.find(qlow)
    if vitri > 350:
        c = r.content[vitri - 350:vitri + 350]
    elif vitri <= 350 and vitri > 0:
        c = rlow[:350]
    else:
        return 0

    a = ['Nội dung vắn tắt', c.decode('utf-8', 'replace'), 'brief_content']
    ketqua.append(a)

    # Cho link avatar pastebin vào [2]
    a = ['Pastebin logo', '/static/images/avatars/pastebin.png', 'pastebin_avatar']
    ketqua.append(a)

    # Đường dẫn URI [3]
    a = ['Đường dẫn URI', b, 'pastebin_uri']
    ketqua.append(a)

    # Loại thông tin [4]
    a = ['Loại thông tin', 'pastebin', 'type_information']
    ketqua.append(a)

    # return
    return ketqua


# Pastebin fetch
def pastebin_fetch(q, first):
    results = []
    ketqua = bing_web_related(q, 'pastebin.com', first)['ketqua']
    for i in ketqua:
        u = urlparse(i)
        b = u.path
        a = pastebin(b, q)
        if a != 0:
            results.append(a)
    return results


# Local data fetch
def local_data_fetch(q, index, page, size):
    results = []
    # Lấy kết quả hồ sơ
    ketqua = full_text_search_es(q, index, page, size)['hits']['hits']
    print ketqua
    for i in ketqua:
        if i['_type'] == "profile":
            a = Report2Present(json.loads(i['_source']['data']))
            if a != 0:
                results.append(a)

    # Lấy kết quả tài khoản
    ketqua = search_by_field('username', q, page, size)['hits']['hits']
    for i in ketqua:
        if i['_type'] == "credential":
            cre_result = []
            cre_result.append(['Tên tài khoản', i['_source']['username'], 'username'])
            cre_result.append(['Mật khẩu', i['_source']['password'], 'password'])
            cre_result.append(['Tiêu đề', i['_source']['title'], 'name'])
            cre_result.append(['Avatar', '/static/images/avatars/cred_local_data.png', 'cred_local_data_avatar'])
            cre_result.append(['Loại dữ liệu', 'cred_local_data', 'type_information'])
            results.append(cre_result)
    return results


# Facebook fanpage
def FbPageExtract(q, cookie):
    results = []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    r = requests.get('https://mbasic.facebook.com/' + q, headers=headers)
    r2 = requests.get('https://m.facebook.com/' + q + '/about', headers=headers)
    content = r.content
    content2 = r2.content
    # Lấy id fanpage
    start = content.find('name="id" value="') + len('name="id" value="')
    end = content.find('"', start)
    fanpage_id = content[start:end]
    results.append(['Fanpage ID', fanpage_id, 'fanpage_id'])
    # Lấy tiếp thông tin đổi tên
    r3 = requests.get('https://mbasic.facebook.com/pages/transparency/' + fanpage_id, headers=headers)
    content3 = r3.content
    # Lấy avatar
    r_avatar = requests.head("https://graph.facebook.com/" + fanpage_id + "/picture?type=large")
    try:
        link_avatar = r_avatar.headers['location']
    except:
        link_avatar = '/static/images/avatars/no_avatar.jpg'
    results.append(['Liên kết hình đại diện', link_avatar, 'fbfp_avatar'])
    # Thêm username
    results.append(['Định danh bằng chữ', q, 'fbfp_username'])
    # Lấy phần giới thiệu
    if content2.find('Giới thiệu</div>') != -1:
        start = content2.find('Giới thiệu</div>') + len('Giới thiệu')
        end = content2.find('id="', start)
        introduce = content2[start:end] + '>'
        introduce = re.sub('<[^>]*>', '', introduce)
        results.append(['Giới thiệu', introduce, 'introduce'])
    else:
        results.append(['Giới thiệu', '', 'introduce'])
    # Lấy Nhiệm vụ
    if content2.find('Nhiệm vụ</div>') != -1:
        start = content2.find('Nhiệm vụ</div>') + len('Nhiệm vụ')
        end = content2.find('id="', start)
        mission = content2[start:end] + '>'
        mission = re.sub('<[^>]*>', '', mission)
        results.append(['Nhiệm vụ', mission, 'mission'])
    else:
        results.append(['Nhiệm vụ', '', 'mission'])
    # Lấy Mô tả
    if content2.find('Mô tả</div>') != -1:
        start = content2.find('Mô tả</div>') + len('Mô tả')
        end = content2.find('id="', start)
        describle = content2[start:end] + '>'
        describle = re.sub('<[^>]*>', '', describle)
        results.append(['Mô tả', describle, 'describle'])
    else:
        results.append(['Mô tả', '', 'describle'])
    # Lấy Tuyên bố quyền sở hữu và quyền tác giả
    if content2.find('Tuyên bố quyền sở hữu và quyền tác giả</div>') != -1:
        start = content2.find('Tuyên bố quyền sở hữu và quyền tác giả</div>') + len(
            'Tuyên bố quyền sở hữu và quyền tác giả')
        end = content2.find('<tbody>', start)
        copyright = content2[start:end]
        copyright = re.sub('<[^>]*>', '', copyright)
        results.append(['Bản quyền', copyright, 'copyright'])
    else:
        results.append(['Bản quyền', '', 'copyright'])
    # Lấy Tiểu sử theo năm
    if content2.find('Tiểu sử theo năm</h3>') != -1:
        start = content2.find('Tiểu sử theo năm</h3>') + len('Tiểu sử theo năm')
        end = content2.find('<a href="/bugnub/?source=FinchPage">', start)
        timeline_content = content2[start:end]
        timeline_content_splited = timeline_content.split('</td>')
        timeline_results = []
        for tcs in xrange(0, len(timeline_content_splited), 1):
            if tcs % 2 == 0:
                time = re.sub('<[^>]*>', '', timeline_content_splited[tcs])
            else:
                event = re.sub('<[^>]*>', '', timeline_content_splited[tcs])
                timeline_results.append({'time': time, 'event': event})
        results.append(['Tiểu sử theo năm', timeline_results, 'timeline'])
    else:
        results.append(['Tiểu sử theo năm', [], 'timeline'])
    # Lấy lịch sử trang
    if content3.find('Lịch sử Trang</div>') != -1:
        start = content3.find('Lịch sử Trang</div>') + len('Lịch sử Trang')
        end = content3.find('Tên thay đổi sẽ cho biết mục đích', start)
        history_content = content3[start:end]
        history_content_splited = history_content.split('font-weight: 400;')[1:]
        history_results = []
        for hcs in xrange(0, len(history_content_splited), 1):
            if hcs % 2 == 0:
                event = re.sub('<[^>]*>', '', '<' + history_content_splited[hcs] + '>')
            else:
                time = re.sub('<[^>]*>', '', '<' + history_content_splited[hcs] + '>')
                history_results.append({'time': time, 'event': event})
        results.append(['Lịch sử trang', history_results, 'history'])
    else:
        results.append(['Lịch sử trang', [], 'history'])
    # Thêm type dữ liệu
    results.append(['Loại dữ liệu', [], 'type_information'])

    return results


# Lấy danh sách
def FetchFbPage(q, page_max, cookie):
    page = 0
    stop = 0
    results = []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:56.0) Gecko/20100101 Firefox/56.0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
               'Accept-Language': 'vi-VN,vi;q=0.8,en-US;q=0.2,en;q=0.2',
               'Referer': 'https://mbasic.facebook.com/search/people/',
               'Origin': 'https://mbasic.facebook.com',
               'Cookie': cookie}
    url = 'https://mbasic.facebook.com/search/pages/?q=' + q
    while page < page_max and stop == 0:
        page += 1
        result = {}
        r = requests.get(url, headers=headers)
        content = r.content
        if content.find('Rất tiếc, chúng tôi không tìm thấy Trang cho') != -1:
            stop = 1
        soup = BeautifulSoup(content, "lxml")
        if len(soup(id="BrowseResultsContainer")) > 0:
            a = str(soup(id="BrowseResultsContainer")[0])
            a = a.split('role="presentation">')
            for i in a[1:]:
                result = {}
                # Lấy tên
                start = i.find('<div class=')
                end = i.find('</div>', start)
                name = i[start:end]
                name = re.sub('<[^>]*>', '', name)
                result['name'] = name
                # Lấy link
                start = i.find('<a href="') + len('<a href="')
                end = i.find('?refid=', start)
                link = i[start:end]
                link = re.sub('<[^>]*>', '', link)
                # Lọc các trường hợp
                link = link.replace('/', '')
                link = link.replace('&amp', '')
                link = link.replace(';', '')
                result['link'] = link
                results.append(result)
        # Sang trang mới
        if content.find('"><span>Xem thêm kết quả</span>') != -1:
            end = content.find('"><span>Xem thêm kết quả</span>')
            start = content.rfind('href="', 0, end) + len('href="')
            new_page_link = content[start:end]
            url = new_page_link
        else:
            stop = 1
    # Lấy thông tin với list đã tìm được
    final_results = []
    for i in results:
        a = FbPageExtract(i['link'], cookie)
        final_results.append(a)
    return final_results


# Main program
# Define web
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = SECRET_KEY
# app update
jinja_options = app.jinja_options.copy()
jinja_options.update(dict(
    variable_start_string='[[',
    variable_end_string=']]'
))
app.jinja_options = jinja_options


# HTTP basic auth


def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'is_login' in session:
            if session['is_login'] == 0:
                return redirect(url_for('login_page'))
        else:
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)

    return decorated_function


# @auth.get_password
# def get_pw(username):
#     users = get_all_users()
#     if username in users:
#         return users.get(username)
#     return None


# Routing index
@app.route('/', methods=['POST', 'GET'])
@app.route('/index', methods=['POST', 'GET'])
@login_required
def index():
    return render_template('index.html')


# API test page
@app.route('/test', methods=['POST', 'GET'])
@login_required
def test():
    return render_template('test.html')


# API cho trang trinh sát tự động
@app.route('/auto_recon_api', methods=['POST', 'GET'])
@login_required
def auto_recon_api():
    results = {}
    results['ketqua'] = []
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        # Tạo mã kiểm tra cho checkbox
        checkbox_check = a['module_selected']
        # assign variable
        if 'email' not in a:
            a['email'] = ''
        if 'phone' not in a:
            a['phone'] = ''
        if 'name' not in a:
            a['name'] = ''
        if 'username' not in a:
            a['username'] = ''
        if 'ckid' not in a:
            a['ckid'] = -1
        # Run functions
        ## Dữ liệu cục bộ, mục to, ưu tiên
        if checkbox_check[5] == "1":
            all_in = a['name'] + '|' + a['phone'] + '|' + a['email'] + '|' + a['username']
            # Phân tách từ khóa tìm kiếm
            splited = all_in.split("|")
            for i in splited:
                if i != '':
                    local_data_result = local_data_fetch(i.encode('utf-8'), 'profiles', 0, 10)
                    # Append vào kết quả chung
                    for j in local_data_result:
                        results['ketqua'].append(j)

        ## username (ưu tiên số I)
        if a['username'] != '':
            fb_results = []
            username_results = {}
            # Kiểm tra showmore
            if 'showmore' not in a:
                a['showmore'] = 0
            # Lay thong tin facebook 1
            if checkbox_check[:2] == "10":
                if a['showmore'] != 1:
                    fb_results = []
                    print apply_cookie(a['ckid'])
                    fb_result = FB(a['username'], apply_cookie(a['ckid']))
                    fb_results.append(fb_result['ketqua'])
                    # Append vào kết quả chung
                    for i in fb_results:
                        results['ketqua'].append(i)

            # FB -> Tim tai khoan twitter 2
            if checkbox_check[:2] == "11" and a['checkbox_ok'] == 1:
                fb_results = []
                fb_results2 = []
                ##look up facebook
                fb_result = FB(a['username'])
                fb_results.append(fb_result['ketqua'])

                ## Find twitter for each fb
                for _fb in fb_results:
                    fb_results2.append(fb_to_tw(_fb))

                # Append vào kết quả chung
                for i in fb_results2:
                    results['ketqua'].append(i)

            # Twitter -> Tim tai khoan facebook 3
            if checkbox_check[:2] == "11" and a['checkbox_ok'] == 2:
                tw_results2 = []

                tw_result = twitter_lookup(a['username'])
                tw_result = twitter_extract(tw_result)
                # Tim facebook
                for _tw in tw_result:
                    tw_results2.append(tw_to_fb(_tw))
                outputed = tw_results2
                # Append vào kết quả chung
                for i in outputed:
                    results['ketqua'].append(i)
            # Lay thong tin twitter 4
            if checkbox_check[:2] == "01":
                if a['showmore'] != 1:
                    tw_result = twitter_lookup(a['username'])
                    tw_result = twitter_extract(tw_result)
                    outputed = tw_result
                    # Append vào kết quả chung
                    for i in outputed:
                        results['ketqua'].append(i)

            # Kiểm tra trên haveibeenpwned 5
            # if a['flag'] != 'showmore':
            #     _haveibeenpwned = haveibeenpwned(a['username'])
            #     for i in _haveibeenpwned:
            #         results['ketqua'].append(i)

        ## phone (ưu tiên số II)
        if a['phone'] != '':
            fb_results = []
            phone_results = {}

            # Lay thong tin facebook 1
            if checkbox_check[0] == "1":
                # Lay id
                fetch_user_FB_results = fetch_user_FB_basic(a['phone'])
                if fetch_user_FB_results == 0:
                    phone_results['ok'] = 0
                    phone_results['ketqua'] = []
                    return json.dumps(phone_results)
                # Lay thong tin facebook roi cho vao list
                for _a in fetch_user_FB_results:
                    fb_result = FB(_a)
                    fb_results.append(fb_result['ketqua'])
                # Append vào kết quả chung
                for i in fb_results:
                    results['ketqua'].append(i)

            # FB -> Tim tai khoan twitter 2
            if checkbox_check[0:1] == "11" and a['checkbox_ok'] == 1:
                fb_results = []
                fb_results2 = []
                # Lay id
                fetch_user_FB_results = fetch_user_FB_basic(a['phone'])
                if fetch_user_FB_results != 0:

                    # Lay thong tin facebook roi cho vao list
                    for _a in fetch_user_FB_results:
                        fb_result = FB(_a)
                        fb_results.append(fb_result['ketqua'])

                    ## Find twitter for each fb
                    for _fb in fb_results:
                        fb_results2.append(fb_to_tw(_fb)[0])

                    # Append vào kết quả chung
                    for i in fb_results2:
                        results['ketqua'].append(i)

            # Twitter -> Tim tai khoan facebook 3 (Bo)
            # if checkbox_check[0:1] == "11" and a['checkbox_ok'] == 2:
            #     tw_results2 = []
            #     phone_results['ok'] = 1
            #     phone_results['ketqua'] = tw_results2

            # Tim tai khoan twitter (Bo) 4
            # if checkbox_check[1] == "1":
            #     fb_results2 = []
            #     phone_results['ok'] = 1
            #     phone_results['ketqua'] = fb_results2

        ## email (Ưu tiên số III)
        if a['email'] != '':
            fb_results = []
            email_results = {}

            # Lay thong tin facebook 1
            if checkbox_check[0] == "1":
                # Lay id
                fetch_user_FB_results = fetch_user_FB_basic(a['email'])
                if fetch_user_FB_results == 0:
                    pass
                # Lay thong tin facebook roi cho vao list
                for _a in fetch_user_FB_results:
                    fb_result = FB(_a)
                    fb_results.append(fb_result['ketqua'])
                # Append vào kết quả chung
                for i in fb_results:
                    results['ketqua'].append(i)

            # FB -> Tim tai khoan twitter 2
            if checkbox_check[0:1] == "11" and a['checkbox_ok'] == 1:
                fb_results = []
                fb_results2 = []
                # Lay id
                fetch_user_FB_results = fetch_user_FB_basic(a['email'])
                if fetch_user_FB_results != 0:
                    # Lay thong tin facebook roi cho vao list
                    for _a in fetch_user_FB_results:
                        fb_result = FB(_a)
                        fb_results.append(fb_result['ketqua'])

                    ## Find twitter for each fb
                    for _fb in fb_results:
                        fb_results2.append(fb_to_tw(_fb)[0])

                    # Append vào kết quả chung
                    for i in fb_results2:
                        results['ketqua'].append(i)

            # Twitter -> Tim tai khoan facebook 3 (Bo)
            if checkbox_check[0:1] == "11" and a['checkbox_ok'] == 2:
                tw_results2 = []
                email_results['ok'] = 1
                email_results['ketqua'] = tw_results2

            # Tim tai khoan twitter (Bo) 4
            if checkbox_check[1] == "1":
                fb_results2 = []
                email_results['ok'] = 1
                email_results['ketqua'] = fb_results2

            # Kiểm tra trên haveibeenpwned 5
            # if a['flag'] != 'showmore':
            #     _haveibeenpwned = haveibeenpwned(a['username'])
            #     for i in _haveibeenpwned:
            #         results['ketqua'].append(i)

        ## name (Ưu tiên số IV)
        if a['name'] != '':
            fb_results = []
            name_results = {}
            if 'page' not in a:
                a['page'] = 1
            # Lay thong tin facebook 1
            if checkbox_check[0] == "1":
                # Lay id
                fetch_user_FB_results = fetch_user_FB(a['name'], a['page'], apply_cookie(a['ckid']))
                # Lay thong tin facebook roi cho vao list
                for _a in fetch_user_FB_results:
                    fb_result = FB(_a['link'])
                    fb_results.append(fb_result['ketqua'])
                # Append vào kết quả chung
                for i in fb_results:
                    results['ketqua'].append(i)

            # FB -> Tim tai khoan twitter 2
            if checkbox_check[0:1] == "11" and a['checkbox_ok'] == 1:
                fb_results = []
                fb_results2 = []
                fetch_user_FB_results = fetch_user_FB(a['name'], a['page'], apply_cookie(a['ckid']))

                # look up facebook
                for _a in fetch_user_FB_results:
                    fb_result = FB(_a)
                    fb_results.append(fb_result['ketqua'])

                ## Find twitter for each fb
                for _fb in fb_results:
                    fb_results2.append(fb_to_tw(_fb))
                # Append vào kết quả chung
                for i in fb_results2:
                    results['ketqua'].append(i)

            # Twitter -> Tim tai khoan facebook 3
            if checkbox_check[0:1] == "11" and a['checkbox_ok'] == 2:
                tw_results2 = []
                name = no_accent_vietnamese(a['name'].encode('utf-8'))
                tw_result = twitter_search(name, a['resultlength'], 5)
                # enum username
                username_string = ''
                for _u in tw_result:
                    username_string += _u['screen_name'] + ','
                username_string = username_string[:-1]  # bo dau phay thua

                tw_result = twitter_lookup(username_string)
                tw_result = twitter_extract(tw_result)
                # Tim facebook
                for _tw in tw_result:
                    tw_results2.append(tw_to_fb(_tw))
                outputed = tw_results2
                # Append vào kết quả chung
                for i in outputed:
                    results['ketqua'].append(i)

            # Lay thong tin twitter 4
            if checkbox_check[1] == "1":
                name = no_accent_vietnamese(a['name'].encode('utf-8'))
                tw_result = twitter_search(name, a['resultlength'], 5)
                # enum username
                username_string = ''
                for _u in tw_result:
                    username_string += _u['screen_name'] + ','
                username_string = username_string[:-1]  # bo dau phay thua

                tw_result = twitter_lookup(username_string)
                tw_result = twitter_extract(tw_result)
                outputed = tw_result
                # Append vào kết quả chung
                for i in outputed:
                    results['ketqua'].append(i)

            # Lay thong tin fanpage 5
            if checkbox_check[6] == "1":
                fbfp_result = FetchFbPage(a['name'], 1, apply_cookie(a['ckid']))
                outputed = fbfp_result
                print outputed
                # Append vào kết quả chung
                for i in outputed:
                    results['ketqua'].append(i)

            # # Lay thong tin tu Bing 6
            # a['checkbox_webs'] = ''
            # if a['checkbox_googleplus'] == False and a['checkbox_facebook'] == False and a[
            #     'checkbox_twitter'] == False and a['checkbox_webs'] == True:
            #     name = no_accent_vietnamese(a['name'].encode('utf-8'))
            #     print name
            #     webs_result = []
            #     name_results['ok'] = 1
            #     name_results['ketqua'] = webs_result

        ## Google plus, mục to
        if checkbox_check[2] == "1":
            if 'bing_page' not in a:
                a['bing_page'] = 0
            all_in = a['name'] + '|' + a['phone'] + '|' + a['email'] + '|' + a['username']
            all_in = no_accent_vietnamese(all_in.encode('utf-8'))
            # Phân tách từ khóa tìm kiếm
            splited = all_in.split("|")
            for i in splited:
                if i != '':
                    gp_result = google_plus_fetch(i, int(a['page']) * 10 + 1)
                    # Append vào kết quả chung
                    for j in gp_result:
                        results['ketqua'].append(j)

        ## Blogger plus, mục to
        if checkbox_check[3] == "1":
            if 'bing_page' not in a:
                a['bing_page'] = 0
            all_in = a['name'] + '|' + a['phone'] + '|' + a['email'] + '|' + a['username']
            all_in = no_accent_vietnamese(all_in.encode('utf-8'))
            # Phân tách từ khóa tìm kiếm
            splited = all_in.split("|")
            for i in splited:
                if i != '':
                    blg_result = blogger_fetch(i, int(a['bing_page']) + 1)
                    # Append vào kết quả chung
                    for j in blg_result:
                        results['ketqua'].append(j)

        ## Pastebin, mục to
        if checkbox_check[4] == "1":
            if 'bing_page' not in a:
                a['bing_page'] = 0
            all_in = a['name'] + '|' + a['phone'] + '|' + a['email'] + '|' + a['username']
            # Phân tách từ khóa tìm kiếm
            splited = all_in.split("|")
            for i in splited:
                if i != '':
                    pastebin_result = pastebin_fetch(i.encode('utf-8'), int(a['bing_page']) + 1)
                    # Append vào kết quả chung
                    for j in pastebin_result:
                        results['ketqua'].append(j)

        # return
        results['ok'] = 1
        return json.dumps(results)
    return 'nothing'


# Đăng nhập (login_page)
@app.route('/login', methods=['POST', 'GET'])
def login_page():
    if request.method == 'GET':
        return render_template('login_page.html')


# Trang tạo chiến dịch mới
@app.route('/new_campaign_page', methods=['POST', 'GET'])
@login_required
def new_campaign_page():
    if request.method == 'POST':
        # Get json
        a = request.get_json()
        b = _full_text_search(a['query'], 0, 5)
        c = search_output(b)
        return json.dumps(c)
    return render_template('new_campaign_page.html')


# Trang chiến dịch của tôi
@app.route('/my_campaign_page', methods=['POST', 'GET'])
@login_required
def my_campaign_page():
    if request.method == 'POST':
        # Get json
        a = request.get_json()
        b = _full_text_search(a['query'], 0, 5)
        c = search_output(b)
        return json.dumps(c)
    return render_template('my_campaign_page.html')


# Trang chiến dịch của tôi
@app.route('/status_analysis_page', methods=['POST', 'GET'])
@login_required
def status_analysis_page():
    if request.method == 'GET':
        try:
            username = request.args.get('username')
            ckid = request.args.get('ckid')
        except(ValueError):
            return 'nothing'
    if username != '':
        return render_template('status_analysis_page.html', username=username, ckid=ckid)
    return 'nothing'


# Trang xử lý báo cáo
@app.route('/report_page', methods=['POST', 'GET'])
@login_required
def report_page():
    if request.method == 'GET':
        try:
            cid = request.args.get('cid')
        except(ValueError):
            return 'nothing'
        if cid != None:
            c = list_saved_record_modal(cid, session['id'])
            content_array = []
            for i in c:
                j = json.loads(i[0])
                j.append(['Saved record ID', i[1], 'rid'])  # Thêm một trường lấy rid để xóa sửa cho dễ
                j.append(['Cookie ID', get_current_cookie_modal(i[2], session['id'])['id'],
                          'ckid'])  # Thêm một trường lấy ckid để báo cáo có thể redirect sang phân tích trạng thái
                j.append(['Campaign ID', i[2], 'cid'])  # Thêm một trường lấy cid để xóa sửa cho dễ
                content_array.append(j)
            results = data_for_stdreport(content_array, session['username'], '', 1)
            return render_template('report_page.html', results=json.dumps(results),
                                   raw_results=json.dumps(content_array))
    return 'nothing'


# Quản lý định danh
@app.route('/manage_cookie_page', methods=['POST', 'GET'])
@login_required
def manage_cookie_page():
    if request.method == 'GET':
        return render_template('manage_cookie_page.html')


# Thêm định danh
@app.route('/add_cookie_page', methods=['POST', 'GET'])
@login_required
def add_cookie_page():
    if request.method == 'GET':
        return render_template('add_cookie_page.html')


# Quản lý dữ liệu nội bộ
@app.route('/manage_local_data_page', methods=['POST', 'GET'])
@login_required
def manage_local_data_page():
    if request.method == 'GET':
        return render_template('manage_local_data_page.html')


# Thêm dữ liệu nội bộ
@app.route('/add_local_data_page', methods=['POST', 'GET'])
@login_required
def add_local_data_page():
    if request.method == 'GET':
        return render_template('add_local_data_page.html')


# Trang quản lý tài khoản
@app.route('/account_page', methods=['POST', 'GET'])
@login_required
def account_page():
    if request.method == 'GET':
        return render_template('account_page.html')


# Help page
@app.route('/help_page', methods=['POST', 'GET'])
@login_required
def help_page():
    return render_template('help_page.html')


# Download file
@app.route('/downloads/<path:filename>', methods=['GET', 'POST'])
def downloadfile(filename):
    a = os.getcwd()
    uploads = app.config['UPLOAD_FOLDER']
    return send_from_directory(directory=uploads, filename=filename)


# Nhận request download file
@app.route('/download_api', methods=['POST', 'GET'])
@login_required
def download_api():
    if request.method == 'POST':
        # Get json
        a = request.get_json()
        if 'profile' in a or 'saved_record_results' in a:
            if a['type'] == 'tech':
                c = download(a['profile'], 'tech', session['username'])
            elif a['type'] == 'std':
                c = download(a['saved_record_results'], 'std', session['username'])
            else:
                return json.dumps({'link': ''})
            return json.dumps({'link': c})
        else:
            return 'nothing'
    return 'nothing'


# Lấy thông tin user hiện tại
@app.route('/get_user_information_api', methods=['POST', 'GET'])
@login_required
def get_user_information_api():
    if request.method == 'POST':
        a = get_information(session['username'])
        return json.dumps(a)
    return 'nothing'


# Nhận username đã được chọn, gộp vào bản ghi
@app.route('/select_username_api', methods=['POST', 'GET'])
@login_required
def select_username_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        profile = a['profile']
        index = a['index']
        if profile[10][2] == 'ratio':
            result = merg_fb_tw(profile[10][1][index][2]['ketqua'], profile)
        else:
            result = merg_fb_tw(profile, profile[27][1][index][2]['ketqua'])
        return json.dumps(result)
    return 'nothing'


# Xem trước API
@app.route('/preview_api', methods=['POST', 'GET'])
@login_required
def preview_api():
    if request.method == 'POST':
        profile = request['profile']
        return json.dumps(profile)
    return 'nothing'


# Kiểm tra internet API
@app.route('/internet_on_api', methods=['POST', 'GET'])
@login_required
def internet_on_api():
    if request.method == 'POST':
        is_internet = internet_on()
        return json.dumps({'internet_on': is_internet})
    return 'nothing'


# Lấy thông tin hệ thống
@app.route('/system_info_api', methods=['POST', 'GET'])
@login_required
def system_info_api():
    if request.method == 'POST':
        res = requests.get('http://just-the-time.appspot.com/')
        time_str = res.content.strip()
        return json.dumps({'ketqua': time_str})
    return 'nothing'


# Tăng lượt truy cập và tìm kiếm
@app.route('/inc_search_api', methods=['POST', 'GET'])
@login_required
def inc_search_api():
    if request.method == 'POST':
        a = ''
        return json.dumps({'ketqua': a})
    return 'nothing'


# Trang tìm kiếm tự động
@app.route('/recon_page', methods=['POST', 'GET'])
@login_required
def recon_page():
    if request.method == 'POST':
        # Get json
        a = request.get_json()
        b = _full_text_search(a['query'], 0, 5)
        c = search_output(b)
        return json.dumps(c)
    if request.method == 'GET':
        try:
            id = int(request.args.get('id'))
        except(ValueError):
            return 'nothing'
        b = get_a_campaign_modal(id, session['id'])
        # Trả hết thông tin
        ## Kiểm tra có dữ liệu không
        if len(b) > 0:
            q = b
            q = json.dumps(q[0])  # Trả về mảng đầu tiên
        else:
            q = []
            q = json.dumps(q)
        return render_template('recon_page.html', q=q)


# Xử lý tạo mới một chiến dịch
@app.route('/new_campaign_api', methods=['POST', 'GET'])
@login_required
def new_campaign_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        # assign variable
        if 'email' not in a:
            a['email'] = ''
        if 'phone' not in a:
            a['phone'] = ''
        if 'name' not in a:
            a['name'] = ''
        if 'username' not in a:
            a['username'] = ''
        if 'checkbox_ok' not in a:
            a['checkbox_ok'] = 0
        if 'ckid' not in a:
            a['ckid'] = -1
        new_campaign_modal(a, session['id'])
        return json.dumps({'ketqua': 'ok'})
    return 'nothing'


# Trả về danh sách chiến dịch
@app.route('/list_campaign_api', methods=['POST', 'GET'])
@login_required
def list_campaign_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        c = list_campaign_modal(session['id'])
        return json.dumps({'ketqua': c})
    return 'nothing'


# Xóa một chiến dịch
@app.route('/delete_campaign_api', methods=['POST', 'GET'])
@login_required
def delete_campaign_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        b = get_information(session['username'])  # Lấy thông tin cơ sở dữ liệu
        try:
            id_campaign = int(a['id'])
        except(ValueError):
            return 'nothing'
        c = delete_campaign_modal(id_campaign, session['id'])
        return json.dumps({'ketqua': 'ok'})
    return 'nothing'


# Sửa một chiến dịch
@app.route('/edit_campaign_api', methods=['POST', 'GET'])
@login_required
def edit_campaign_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        b = get_information(session['username'])  # Lấy thông tin cơ sở dữ liệu

        c = edit_campaign_modal(a, session['id'])
        return json.dumps({'ketqua': 'ok'})
    return 'nothing'


## Quản lý các bản ghi đã lưu

# Liệt kê các bản ghi đã lưu cho 1 campaign
@app.route('/list_saved_record_api', methods=['POST', 'GET'])
@login_required
def list_saved_record_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        c = list_saved_record_modal(a['cid'], session['id'])
        content_array = []
        for i in c:
            j = json.loads(i[0])
            j.append(['Saved record ID', i[1], 'rid'])  # Thêm một trường lấy rid để xóa sửa cho dễ
            j.append(['Campaign ID', i[2], 'cid'])  # Thêm một trường lấy cid để xóa sửa cho dễ
            content_array.append(j)
        # Lấy ckid
        current_ckid = get_current_cookie_modal(a['cid'], session['id'])['id']
        return json.dumps({'ketqua': content_array, 'ckid': current_ckid})
    return 'nothing'


# Lưu một bản ghi mới
@app.route('/new_saved_record_api', methods=['POST', 'GET'])
@login_required
def new_saved_record_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        new_saved_record_modal(a, session['id'])
        return json.dumps({'ketqua': 'ok'})
    return 'nothing'


# Xóa một bản ghi đã lưu
@app.route('/delete_saved_record_api', methods=['POST', 'GET'])
@login_required
def delete_saved_record_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        b = get_information(session['username'])  # Lấy thông tin cơ sở dữ liệu
        try:
            rid = int(a['rid'])
        except(ValueError):
            return 'nothing'
        c = delete_saved_record_modal(rid, session['id'])
        return json.dumps({'ketqua': 'ok'})
    return 'nothing'


# Liẹt kê status
@app.route('/list_status_api', methods=['POST', 'GET'])
@login_required
def list_status_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        PageNumber = int(a['PageNumber'])
        SearchText = no_accent_vietnamese(a['searchtext'])
        if 'username' in a and 'ckid' in a:
            if a['ckid'] == 'None':
                ckid = -1
            else:
                ckid = int(a['ckid'])
            results = GetFbStatus(a['username'], apply_cookie(ckid), PageNumber, 1, 1, a['NextURL'])
            # Tìm kiếm
            results['results'] = SearchFbStatus(results['results'], SearchText, 'value', 'comments')
        else:
            results = {'results': [], 'NextURL': ''}
        return json.dumps({'ketqua': results})
    return 'nothing'


## API liên quan tới quản lý cookie
# Trả về danh sách chiến dịch
@app.route('/list_cookie_api', methods=['POST', 'GET'])
@login_required
def list_cookie_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        c = list_cookies_modal(session['id'])
        return json.dumps({'ketqua': c})
    return 'nothing'


# Xóa một chiến dịch
@app.route('/delete_cookie_api', methods=['POST', 'GET'])
@login_required
def delete_cookie_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        try:
            ckid = int(a['ckid'])
        except(ValueError):
            return 'nothing'
        c = delete_cookies_modal(ckid, session['id'])
        return json.dumps({'ketqua': 'ok'})
    return 'nothing'


# Sửa một chiến dịch
@app.route('/edit_cookie_api', methods=['POST', 'GET'])
@login_required
def edit_cookie_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json

        c = edit_cookies_modal(a['name'], a['cookie'], a['description'], a['ckid'], session['id'])
        return json.dumps({'ketqua': 'ok'})
    return 'nothing'


# Xử lý tạo mới một cookie
@app.route('/new_cookie_api', methods=['POST', 'GET'])
@login_required
def new_cookie_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        # assign variable
        if 'name' not in a:
            a['name'] = ''
        if 'value' not in a:
            a['value'] = ''
        if 'description' not in a:
            a['description'] = ''
        new_cookies_modal(a, session['id'])
        return json.dumps({'ketqua': 'ok'})
    return 'nothing'


# Xử lý tạo mới một cookie
@app.route('/get_current_cookie_api', methods=['POST', 'GET'])
@login_required
def get_current_cookie_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        # assign variable
        if 'cid' not in a:
            a['cid'] = ''
        # Lấy định hiện tại mà chiến dịch sử dụng
        current_cookie = get_current_cookie_modal(a['cid'], session['id'])
        # Lấy danh sách định danh của người dùng
        result = list_cookies_modal(session['id'])
        # Tạo định dạng json chuẩn để gửi cho bootstrap checkbox
        results = {}
        results['availableOptions'] = []
        results['selectedOption'] = {}
        for i in result:
            results['availableOptions'].append({'id': i[0], 'name': i[1]})
        # Thêm định danh mặc định
        results['availableOptions'].append({'id': -1, 'name': 'Mặc định'})
        results['selectedOption'] = current_cookie
        return json.dumps(results)
    return 'nothing'


## API liên quan tới elasticsearch
# <profile>
# Liệt kê các bản ghi profile
@app.route('/list_profile_local_data_api', methods=['POST', 'GET'])
@login_required
def list_profile_local_data_api():
    if request.method == 'POST':
        b = request.get_json()  # Receive data as json
        c = list_profile(session['id'], 'profiles', 'profile')
        # return json.dumps({'ketqua': c})
        return json.dumps({'data': c})
    return 'nothing'


# Lưu 1 bản ghi profile vào elasticsearch
@app.route('/save_es_api', methods=['POST', 'GET'])
@login_required
def save_es_api():
    if request.method == 'POST':
        b = request.get_json()  # Receive data as json
        result = add_document_es(b['q'], 'profile', session['id'], 'profiles')
        return 'ok'
    return 'nothing'


# Xóa 1 bản ghi profile elasticsearch
@app.route('/delete_profile_local_data_api', methods=['POST', 'GET'])
@login_required
def delete_profile_local_data_api():
    if request.method == 'POST':
        b = request.get_json()  # Receive data as json
        result = delete_document_es(b['id'], 'profiles', b['type'])
        return 'ok'
    return 'nothing'


# <credential>
# Xử lý file để import cho hai loại: tài khoản
@app.route('/file_import_es_api', methods=['POST', 'GET'])
@login_required
def file_import_es_api():
    if request.method == 'POST':
        filestorage = request.files['file']
        a = filestorage.read()
        title = request.form['title']
        if title != '':
            result = checkFormat(a, title, 'profiles', 'credential', session['id'])
            return json.dumps({'content': 'ok'})
        else:
            return json.dumps({'content': 'nothing'})


# Đếm số bản ghi cho từng đợt tải lên
@app.route('/list_cred_api', methods=['POST', 'GET'])
@login_required
def list_cred_api():
    if request.method == 'POST':
        if request.method == 'POST':
            b = request.get_json()  # Receive data as json
            result = list_cred(session['id'], 'profiles', 'credential')
            return json.dumps({'result': result})
        return 'nothing'


# Xóa 1 bản ghi cred elasticsearch
@app.route('/delete_cred_local_data_api', methods=['POST', 'GET'])
@login_required
def delete_cred_local_data_api():
    if request.method == 'POST':
        b = request.get_json()  # Receive data as json
        result = delete_cred_local_data_es(b['title'], b['timeStamp'], session['id'])
        return 'ok'
    return 'nothing'


## API liên quan tới tài khoản
# Thay đổi mật khẩu
@app.route('/change_password_api', methods=['POST', 'GET'])
@login_required
def change_password_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        id = session['id']
        password = a['new_password']
        old_password = a['old_password']
        b = change_password(id, password, old_password)
        if b == 1:
            return json.dumps({'is_ok': 'ok'})
        elif b == 0:
            return json.dumps({'is_ok': 'NotMatchPW'})
    return 'nothing'


# Đăng nhập
@app.route('/login_api', methods=['POST', 'GET'])
def login_api():
    if request.method == 'POST':
        a = request.get_json()  # Receive data as json
        username = a['username']
        password = a['password']
        if verify_login(username, password) == 1:
            session['is_login'] = 1
            session['username'] = username
            user_information = get_information(username)
            session['id'] = user_information[len(user_information) - 1]
            return json.dumps({'is_ok': 'ok'})
        return json.dumps({'is_ok': 'fail'})
    return 'nothing'


# Đăng xuất
@app.route('/logout_api', methods=['POST', 'GET'])
def logout_api():
    if request.method == 'POST':
        session['is_login'] = 0
        session['username'] = ''
        session['id'] = -1
        return json.dumps({'is_ok': 'ok'})
    return 'nothing'


# App run
app.run(debug=True, host='0.0.0.0', port=5000, threaded=True, ssl_context='adhoc')
SESSION_TYPE = 'redis'
app.config.from_object(__name__)
Session(app)
